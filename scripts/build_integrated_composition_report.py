#!/usr/bin/env python3
"""Build an integrated learner MF/MD and native-reference analysis report."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from .build_mfmd_report import (
        BLUE,
        DARK_BLUE,
        GOLD,
        INK,
        LIGHT_FILL,
        MUTED,
        add_callout,
        add_dataframe_table,
        add_field,
        add_figure,
        add_paragraph,
        add_table_source,
        configure_document,
        set_run_font,
    )
except ImportError:
    from build_mfmd_report import (
        BLUE,
        DARK_BLUE,
        GOLD,
        INK,
        LIGHT_FILL,
        MUTED,
        add_callout,
        add_dataframe_table,
        add_field,
        add_figure,
        add_paragraph,
        add_table_source,
        configure_document,
        set_run_font,
    )


DEFAULT_LEARNER_ANALYSIS_DIR = "outputs/mfmd_analysis"
DEFAULT_NATIVE_ANALYSIS_DIR = "outputs/native_control_analysis"
DEFAULT_LEARNER_WORKBOOK = "作文词性统计宽表.xlsx"
DEFAULT_NATIVE_MASTER = "母语作文样本主表.xlsx"
DEFAULT_NATIVE_WORKBOOK = "母语作文词性统计宽表.xlsx"
DEFAULT_OUTPUT = "作文语言特征与母语参照综合分析报告.docx"

DIMENSIONS = [
    "词汇丰富度与词汇扩展",
    "基础词汇与叙事推进",
    "人称指涉与信息密度",
    "句法延展与分句复杂度",
    "动作过程与动词链",
]

DIMENSION_INTERPRETATIONS = {
    "词汇丰富度与词汇扩展": (
        "该维度集中反映词形多样性、篇内低重复程度以及名词和动词词库的扩展。"
        "高分表示同等篇幅内动用的不同词形较多，不等同于词语更生僻，也不能单独代表作文质量。"
    ),
    "基础词汇与叙事推进": (
        "该维度的正端由时间词、趋向动词、低等级词汇覆盖和专名推进，负端更多关联中高等级词汇及结构助词。"
        "它描述叙事推进所依赖的基础词汇资源与较抽象表达之间的相对配置。"
    ),
    "人称指涉与信息密度": (
        "该维度把代词和人称参与置于正端，把词汇密度、平均词长及部分非HSK词汇置于另一端。"
        "它主要区分参与者显性指涉与名词化、信息压缩表达，而不是口语和书面语的简单优劣。"
    ),
    "句法延展与分句复杂度": (
        "该维度由平均句长、句长变异、长句、逗号和分句数共同定义，反映句子内部展开与从句组织。"
        "网页作文可能经过修改或编辑，因此母语参照在该维度上的优势尤其需要结合来源条件解释。"
    ),
    "动作过程与动词链": (
        "该维度由动词频率、情态动词与连续动词结构构成，反映事件链、动作过程及行动可能性的表达。"
        "记叙题通常比议论题更需要动作链，跨题目比较时必须保留任务功能这一前提。"
    ),
}

PAIR_LABELS = {
    "J1": "人物影响记叙",
    "J2": "假期旅行及个人经历记叙",
    "Y1": "健康与相邻社会议题",
    "Y2": "父母教育与成长议题",
}

# The page numbers are updated after the final render QA pass.
TOC_ENTRIES = [
    ("摘要", 3),
    ("1 研究背景与整合问题", 4),
    ("2 数据来源与800篇样本", 6),
    ("3 301项语言特征与统计方法", 9),
    ("4 学习者作文的五维结构", 12),
    ("5 学习者四组、分数与国籍", 16),
    ("6 公开网络母语参照样本审计", 23),
    ("7 四组母语参照对照", 26),
    ("8 稳健性与联合因子分析", 33),
    ("9 五维综合证据", 37),
    ("10 教学与研究启示", 40),
    ("11 局限", 40),
    ("12 结论", 41),
    ("参考文献与网页来源", 41),
    ("附录A 最终因子载荷", 43),
    ("附录B 母语样本审计摘要", 44),
    ("附录C 可复现性说明", 44),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成作文语言特征与母语参照综合分析报告。")
    parser.add_argument("--learner-analysis-dir", default=DEFAULT_LEARNER_ANALYSIS_DIR)
    parser.add_argument("--native-analysis-dir", default=DEFAULT_NATIVE_ANALYSIS_DIR)
    parser.add_argument("--learner-workbook", default=DEFAULT_LEARNER_WORKBOOK)
    parser.add_argument("--native-master", default=DEFAULT_NATIVE_MASTER)
    parser.add_argument("--native-workbook", default=DEFAULT_NATIVE_WORKBOOK)
    parser.add_argument("--output", default=DEFAULT_OUTPUT)
    return parser.parse_args()


def read_table(directory: Path, name: str) -> pd.DataFrame:
    path = directory / "tables" / f"{name}.csv"
    if not path.is_file():
        raise FileNotFoundError(f"找不到分析表：{path}")
    return pd.read_csv(path, encoding="utf-8-sig")


def p_text(value: float) -> str:
    return "p<0.001" if value < 0.001 else f"p={value:.3f}"


def effect_label(value: float) -> str:
    magnitude = abs(value)
    if magnitude < 0.20:
        return "可忽略"
    if magnitude < 0.50:
        return "小"
    if magnitude < 0.80:
        return "中等"
    return "大"


def compact_native_regression_term(value: Any) -> str:
    term = str(value)
    if ":C(对应题目" in term:
        for code in ("J2", "Y1", "Y2"):
            if f"[T.{code}]" in term:
                return f"语料来源×{code}"
    if "公开网络母语参照" in term:
        return "语料来源：母语参照（J1）"
    return term


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.208
        set_run_font(paragraph.add_run(item))


def configure_integrated_header(document: Document) -> None:
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "作文语言特征与母语参照综合分析"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(header.runs[0], size=8.5, color=MUTED)
        footer = section.footer.paragraphs[0]
        if not footer.text:
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_run_font(footer.add_run("第 "), size=8.5, color=MUTED)
            add_field(footer, "PAGE")
            set_run_font(footer.add_run(" 页"), size=8.5, color=MUTED)


def add_cover(document: Document, learner_meta: dict[str, Any], native_meta: dict[str, Any]) -> None:
    for _ in range(5):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("语料库多维统计综合研究报告"), size=11, bold=True, color=GOLD)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("作文语言特征与母语参照\n综合分析"), size=28, bold=True, color=INK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    learner_rows = learner_meta["validation"]["rows"]
    native_rows = native_meta["validation"]["native_rows"]
    set_run_font(
        subtitle.add_run(
            f"基于{learner_rows}篇HSK学习者作文与{native_rows}篇公开网络高中作文的MF/MD研究"
        ),
        size=13.5,
        color=DARK_BLUE,
    )

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(28)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    rule._p.get_or_add_pPr().append(borders)

    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        summary.add_run(
            f"800篇作文 · 301列统计 · {learner_meta['selected_model']['feature_count']}项最终指标 · "
            f"{learner_meta['selected_model']['factor_count']}个语言维度"
        ),
        size=10.5,
        bold=True,
        color=INK,
    )
    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.paragraph_format.space_before = Pt(64)
    set_run_font(date_line.add_run(f"{date.today().year}年{date.today().month}月"), size=11, color=MUTED)
    document.add_page_break()


def add_static_toc(document: Document) -> None:
    document.add_heading("目录", level=1)
    frame = pd.DataFrame(TOC_ENTRIES, columns=["章节", "页码"])
    add_dataframe_table(
        document,
        frame,
        columns=["章节", "页码"],
        widths_dxa=[8200, 1160],
        font_size=8.8,
    )
    add_table_source(document, "目录页码按最终渲染版本回填。")
    document.add_page_break()


def validate_inputs(
    learner_meta: dict[str, Any],
    native_meta: dict[str, Any],
    learner_stats: pd.DataFrame,
    native_stats: pd.DataFrame,
    native_master: pd.DataFrame,
    assignments: pd.DataFrame,
    native_comparisons: pd.DataFrame,
) -> None:
    learner_validation = learner_meta["validation"]
    native_validation = native_meta["validation"]
    if learner_validation["rows"] != 620 or len(learner_stats) != 620:
        raise ValueError("学习者作文必须为620篇")
    if native_validation["native_rows"] != 180 or len(native_stats) != 180 or len(native_master) != 180:
        raise ValueError("母语参照作文必须为180篇")
    if learner_stats.shape[1] != 301 or native_stats.shape[1] != 301:
        raise ValueError("两份词性统计宽表都必须为301列")
    if learner_validation["code_counts"] != {"J1": 155, "J2": 155, "Y1": 155, "Y2": 155}:
        raise ValueError("学习者四组样本量不是各155篇")
    if native_validation["native_group_counts"] != {"NJ1": 45, "NJ2": 45, "NY1": 45, "NY2": 45}:
        raise ValueError("母语参照四组样本量不是各45篇")
    if len(assignments) != 39 or native_validation["selected_features"] != 39:
        raise ValueError("五维投影必须使用39项最终指标")
    if native_comparisons.shape[0] != 20:
        raise ValueError("母语五维比较必须为4组×5维共20项")
    if int((native_comparisons["Holm校正p值"] < 0.05).sum()) != 13:
        raise ValueError("母语比较显著项与正式结果不一致")
    if set(native_master["学段"]) != {"高中"}:
        raise ValueError("本轮母语参照样本应全部为高中范围")
    if native_master["来源URL"].duplicated().any() or native_master["正文SHA256"].duplicated().any():
        raise ValueError("母语主表存在重复URL或正文哈希")


def learner_dimension_table(
    descriptives: pd.DataFrame,
    welch: pd.DataFrame,
) -> pd.DataFrame:
    pivot = descriptives.pivot(index="维度名称", columns="篇名代码", values="均值").reset_index()
    result = pivot.merge(
        welch[["维度名称", "Omega平方", "BH校正p值"]],
        on="维度名称",
        how="left",
    )
    return result.rename(columns={"维度名称": "维度"})


def hsk_summary(learner_stats: pd.DataFrame, native_stats: pd.DataFrame) -> pd.DataFrame:
    columns = ["初等词汇占比", "中等词汇占比", "高等词汇占比", "非HSK词汇占比", "HSK词汇覆盖率"]
    learner = learner_stats.groupby("篇名代码", as_index=False)[columns].mean().rename(columns={"篇名代码": "代码"})
    native = native_stats.groupby("篇名代码", as_index=False)[columns].mean().rename(columns={"篇名代码": "代码"})
    result = pd.concat([learner, native], ignore_index=True)
    order = ["J1", "NJ1", "J2", "NJ2", "Y1", "NY1", "Y2", "NY2"]
    result["_order"] = result["代码"].map({value: index for index, value in enumerate(order)})
    return result.sort_values("_order").drop(columns="_order")


def source_summary(native_master: pd.DataFrame) -> pd.DataFrame:
    result = native_master.groupby("来源类型", as_index=False).size().rename(columns={"size": "篇数"})
    result["占比"] = result["篇数"] / len(native_master)
    return result.sort_values(["篇数", "来源类型"], ascending=[False, True]).reset_index(drop=True)


def year_summary(native_master: pd.DataFrame) -> pd.DataFrame:
    data = native_master.copy()
    data["年份"] = pd.to_datetime(data["发布日期"]).dt.year
    result = (
        data.groupby("母语代码", as_index=False)
        .agg(
            样本量=("作文文件名", "size"),
            最早年份=("年份", "min"),
            最晚年份=("年份", "max"),
            平均汉字数=("正文汉字数", "mean"),
            平均目标偏差=("目标篇幅相对偏差", "mean"),
            年代扩展样本=("年代扩展样本", "sum"),
        )
    )
    return result


def synthesis_table(
    learner_descriptives: pd.DataFrame,
    score_correlations: pd.DataFrame,
    native_comparisons: pd.DataFrame,
    resampling: pd.DataFrame,
) -> pd.DataFrame:
    records: list[dict[str, Any]] = []
    for dimension in DIMENSIONS:
        learner = learner_descriptives.loc[learner_descriptives["维度名称"] == dimension]
        high = learner.loc[learner["均值"].idxmax()]
        low = learner.loc[learner["均值"].idxmin()]
        score = score_correlations.loc[score_correlations["维度名称"] == dimension].iloc[0]
        native = native_comparisons.loc[native_comparisons["维度"] == dimension].copy()
        strongest = native.loc[native["Hedges_g_母语减学习者"].abs().idxmax()]
        stable = resampling.loc[resampling["维度"] == dimension]
        stable_count = int(
            ((stable["重抽样95%CI下限"] > 0) | (stable["重抽样95%CI上限"] < 0)).sum()
        )
        records.append(
            {
                "维度": dimension,
                "学习者最高/最低": f"{high['篇名代码']} / {low['篇名代码']}",
                "分数rho": score["Spearman_rho"],
                "最强母语对照": f"{strongest['对应题目']} g={strongest['Hedges_g_母语减学习者']:.2f}",
                "Holm显著组数": int((native["Holm校正p值"] < 0.05).sum()),
                "篇幅重抽样稳定组数": stable_count,
            }
        )
    return pd.DataFrame(records)


def build_report(
    learner_analysis_dir: Path,
    native_analysis_dir: Path,
    learner_workbook: Path,
    native_master_path: Path,
    native_workbook: Path,
    output_path: Path,
) -> None:
    learner_meta = json.loads((learner_analysis_dir / "analysis_metadata.json").read_text(encoding="utf-8"))
    native_meta = json.loads((native_analysis_dir / "analysis_metadata.json").read_text(encoding="utf-8"))
    learner_stats = pd.read_excel(learner_workbook, sheet_name="词性统计")
    native_stats = pd.read_excel(native_workbook, sheet_name="词性统计")
    native_master = pd.read_excel(native_master_path, sheet_name="母语样本主表")

    learner_sample = read_table(learner_analysis_dir, "样本概况")
    screening = read_table(learner_analysis_dir, "变量筛选")
    model_diagnostics = read_table(learner_analysis_dir, "候选模型诊断")
    loadings = read_table(learner_analysis_dir, "因子载荷")
    assignments = read_table(learner_analysis_dir, "指标维度归属")
    learner_descriptives = read_table(learner_analysis_dir, "维度描述统计")
    learner_welch = read_table(learner_analysis_dir, "Welch检验")
    score_correlations = read_table(learner_analysis_dir, "分数相关")
    learner_regressions = read_table(learner_analysis_dir, "稳健回归")

    native_sampling = read_table(native_analysis_dir, "采样审计概况")
    native_descriptives = read_table(native_analysis_dir, "维度描述统计")
    native_comparisons = read_table(native_analysis_dir, "五维Welch比较")
    native_features = read_table(native_analysis_dir, "39项特征比较")
    native_regressions = read_table(native_analysis_dir, "HC3稳健回归")
    resampling = read_table(native_analysis_dir, "篇幅匹配重抽样")
    sensitivity = read_table(native_analysis_dir, "主题年份敏感性")
    joint_models = read_table(native_analysis_dir, "联合模型诊断")
    congruence = read_table(native_analysis_dir, "Tucker一致性")

    validate_inputs(
        learner_meta,
        native_meta,
        learner_stats,
        native_stats,
        native_master,
        assignments,
        native_comparisons,
    )

    learner_figures = learner_analysis_dir / "figures"
    native_figures = native_analysis_dir / "figures"
    figure_paths = [
        learner_figures / "01_样本与分数结构.png",
        native_figures / "01_采样主题层级.png",
        learner_figures / "02_变量筛选流程.png",
        learner_figures / "03_平行分析与碎石图.png",
        learner_figures / "04_因子载荷热图.png",
        learner_figures / "05_四组维度得分分布.png",
        learner_figures / "06_四组多维画像热图.png",
        learner_figures / "07_分数与维度得分趋势.png",
        learner_figures / "08_HSK等级构成.png",
        learner_figures / "09_组间效应量森林图.png",
        learner_figures / "10_回归系数森林图.png",
        learner_figures / "11_维度相关热图.png",
        native_figures / "02_篇幅匹配诊断.png",
        native_figures / "03_五维画像热图.png",
        native_figures / "05_五维效应量森林图.png",
        native_figures / "06_主要语言特征效应量.png",
        native_figures / "07_HSK词汇构成.png",
        native_figures / "08_篇幅匹配重抽样.png",
        native_figures / "10_联合模型Tucker一致性.png",
    ]
    missing_figures = [str(path) for path in figure_paths if not path.is_file()]
    if missing_figures:
        raise FileNotFoundError(f"综合报告缺少图片：{missing_figures}")

    selected_model = model_diagnostics.loc[model_diagnostics["是否选定"]].copy()
    if len(selected_model) != 1:
        raise ValueError("学习者候选模型中必须且只能选定一个模型")
    selected_model = selected_model.iloc[0]
    native_significant = int((native_comparisons["Holm校正p值"] < 0.05).sum())
    resampling_robust = int(
        ((resampling["重抽样95%CI下限"] > 0) | (resampling["重抽样95%CI上限"] < 0)).sum()
    )
    score_significant = int((score_correlations["BH校正p值"] < 0.05).sum())

    document = Document()
    configure_document(document)
    configure_integrated_header(document)
    add_cover(document, learner_meta, native_meta)
    add_static_toc(document)

    document.add_heading("摘要", level=1)
    add_paragraph(
        document,
        "本报告把项目既有的学习者作文MF/MD分析和公开网络母语参照分析整合为一条连续证据链。"
        "学习者语料包括J1、J2、Y1、Y2四组各155篇，共620篇；母语参照包括NJ1、NJ2、NY1、NY2四组各45篇，共180篇。"
        "两类语料均经过相同的PyNLPIR分词与301列语言特征统计，总样本量为800篇。母语参照来自作文网高中作文栏目，"
        "可以证明网页栏目归类，但不能独立验证作者身份或排除编辑、转载和润色，因此全文统一使用“公开网络母语参照语料”这一审慎名称。",
    )
    add_paragraph(
        document,
        f"学习者分析从146项直接指标出发，经方差、稀疏性、相关性和KMO筛选，并使用MinRes提取、Promax旋转、"
        f"平行分析与Bootstrap稳定性检验，最终得到5个维度和39项指标。所选模型KMO={selected_model['KMO']:.3f}，"
        f"累计解释方差为{selected_model['累计解释方差']:.1%}。母语分析固定使用学习者样本的均值、标准差和载荷方向投影，"
        f"避免用母语样本重新定义量尺。20项题目配对维度比较中有{native_significant}项经Holm校正后显著，"
        f"{resampling_robust}项在1000次篇幅匹配重抽样中的95%区间不跨0。",
    )
    add_paragraph(
        document,
        f"学习者内部分析显示，题目与体裁组合在五个维度上均存在结构性差异，{score_significant}个维度与作文分数呈校正后显著相关。"
        "母语参照进一步揭示，学习者与公开网络高中作文的差异不是单一的“高级或低级”，而是分布在词汇丰富度、基础词汇配置、"
        "人称指涉、句法延展和动作链等多个方向。联合800篇作文重新探索因子结构时，只有2因子方案同时通过诊断与稳定性检验，"
        "说明原五维结构适合作为固定比较量尺，但不能被视为跨来源完全不变的普遍结构。",
    )
    add_callout(
        document,
        "核心结论",
        "五维得分描述的是语言资源配置，不是作文质量排名。最可信的结论来自题目配对、效应量、稳健回归、篇幅重抽样和主题敏感性共同支持的方向；"
        "母语参照结果应解释为语料来源差异，而不是经过身份认证的母语能力因果效应。",
        fill="FFF7E8",
    )
    add_paragraph(document, "关键词：HSK作文；MF/MD；语言特征；母语参照；因子分析；词汇丰富度；句法复杂度")

    document.add_page_break()
    document.add_heading("1 研究背景与整合问题", level=1)
    add_paragraph(
        document,
        "学习者作文研究经常依赖字数、词汇等级或单项语法指标，但这些指标彼此相关，单独解释容易把题目需求、篇幅和语言能力混为一谈。"
        "Biber提出的多特征多维度方法强调先观察多项语言特征的共同变异，再以因子载荷解释潜在功能维度。徐勤2021年的叙述文研究和2023年博士论文"
        "把这一思路用于日本人汉语学习者与汉语母语者作文比较，为本项目提供了指标组织、斜交旋转和维度解释的直接方法参照。",
    )
    add_paragraph(
        document,
        "本项目首先在620篇HSK作文内部建立五维量尺，回答不同题目、分数及补充国籍因素如何与语言资源配置相关。随后，项目从公开网络高中作文中"
        "构建180篇参照语料，并将其投影到完全相同的五维量尺。两阶段分析原本分别报告，容易造成方法说明重复，也不利于判断某个维度究竟是学习者内部差异、"
        "分数差异，还是学习者与参照语料之间的来源差异。本综合报告因此以同一组研究问题串联两阶段证据。",
    )
    add_callout(
        document,
        "整合研究问题",
        "RQ1：620篇学习者作文呈现怎样的稳定多维结构？RQ2：J1、J2、Y1、Y2在五维上有何差异？RQ3：作文分数及日本籍控制变量与维度得分有何关系？"
        "RQ4：四组公开网络母语参照与对应学习者题目在五维和具体指标上有何差异？RQ5：这些差异能否经受篇幅、主题、年份和联合因子结构检验？",
    )
    document.add_heading("1.1 两阶段证据的关系", level=2)
    add_paragraph(
        document,
        "学习者阶段的五维模型是主量尺，母语参照阶段是外部投影。投影过程固定39项指标的标准化参数及载荷正负方向，因此母语样本不会反向改变原模型。"
        "这一设计优于把800篇作文直接混合后重新命名维度，因为后者无法区分原学习者结构和新增语料引起的结构变化。联合EFA只作为敏感性分析，专门检查这种变化。",
    )
    document.add_heading("1.2 解释边界", level=2)
    add_bullets(
        document,
        [
            "J1/J2与Y1/Y2同时包含题目和体裁差异，不能将组间差异解释为纯粹体裁效应。",
            "公开网页高中作文不是身份、年龄和写作条件均可核验的实验母语语料，不能直接推出母语能力因果结论。",
            "五维正负方向来自统计载荷，不表示好坏；维度得分高只表示相应语言资源更集中。",
            "作文分数缺少75分层级，并且是既有评分结果；相关与回归分析不构成评分因果模型。",
        ],
    )
    evidence_levels = pd.DataFrame(
        [
            ["描述层", "各组使用了哪些语言资源", "均值、比例、每千字频率", "不能说明差异是否稳定"],
            ["学习者推断层", "题目组和分数是否相关", "Welch、效应量、Spearman、HC3", "不能推出体裁或分数的因果效应"],
            ["外部参照层", "同题参照与学习者差在哪里", "固定投影、Hedges g、Bootstrap", "不能认证网页作者身份"],
            ["稳健性层", "结论是否依赖篇幅、主题或模型", "重抽样、敏感性、联合EFA", "不能消除所有未观测混杂"],
        ],
        columns=["证据层级", "回答问题", "主要方法", "解释边界"],
    )
    add_dataframe_table(
        document,
        evidence_levels,
        columns=["证据层级", "回答问题", "主要方法", "解释边界"],
        widths_dxa=[1300, 2400, 2500, 3160],
        font_size=8.0,
    )
    add_table_source(document, "表0  综合报告的四层证据结构。")

    document.add_page_break()
    document.add_heading("2 数据来源与800篇样本", level=1)
    document.add_heading("2.1 学习者作文", level=2)
    add_paragraph(
        document,
        "学习者语料来自HSK动态作文语料库筛选结果。J1为人物影响记叙，J2为假期经历记叙，Y1为吸烟、健康与公共利益议论，Y2为父母教育与儿童成长议论。"
        "抽样在各分数层内保持四组数量一致，并尽量扩大国籍覆盖、平衡同一国籍数量；排除中国香港、中国台湾、中国澳门和中国少数民族样本。"
        "75分样本已从主表移除，最终每组155篇。",
    )
    add_dataframe_table(
        document,
        learner_sample,
        columns=["篇名代码", "篇名", "体裁", "样本量", "平均分", "平均纯文本字数", "国籍数"],
        widths_dxa=[700, 2600, 900, 800, 850, 1700, 1810],
        font_size=8.0,
    )
    add_table_source(document, "表1  学习者作文样本结构。每组155篇，总计620篇。")
    add_figure(document, learner_figures / "01_样本与分数结构.png", 1, "学习者作文的样本、体裁与分数结构")

    document.add_heading("2.2 公开网络母语参照作文", level=2)
    add_paragraph(
        document,
        "参照候选来自作文网高中作文栏目、高一至高三常规作文页面及高中作文文库。完整高中单元作文和练习来源作文可以纳入；"
        "仅有题目、要求、讲解或素材的练习页仍排除。另排除满分作文、范文、写作指导、点评、解析、读后感、演讲稿、小说和诗歌。"
        "NY2保留了初中作文最后备用通道，但最终180篇全部来自高中范围，未启用初中备选。",
    )
    add_paragraph(
        document,
        "目标原为每组50篇。完成候选发现、正文提取、篇幅约束、体裁校验、正文哈希与字符五元组近重复去重后，NJ2只有45篇合格候选，"
        "因此四组统一缩减为各45篇。主题匹配分为精确、近似、扩展和宽泛四级；为保持高中范围，NY1、NY2及部分NJ2放宽到相邻题目域。"
        "这些宽泛主题提高了样本量，也降低了严格同题可比性。",
    )
    add_dataframe_table(
        document,
        native_sampling,
        columns=["母语代码", "样本量", "平均汉字数", "精确主题", "近似主题", "扩展主题", "宽泛主题", "高一", "高二", "高三"],
        widths_dxa=[700, 700, 1050, 850, 850, 850, 850, 1170, 1170, 1170],
        font_size=7.5,
    )
    add_table_source(document, "表2  母语参照样本的主题匹配和年级结构。")
    add_figure(document, native_figures / "01_采样主题层级.png", 2, "四组母语参照样本的主题匹配层级")

    document.add_heading("2.3 文本处理与版权边界", level=2)
    add_paragraph(
        document,
        "两类语料都保留原文、清洗正文和PyNLPIR分词文本三个层次。学习者原文中的网站标注按可确定规则清理；母语网页正文仅删除标题、日期、来源、广告、"
        "二维码和版权尾注，不修改措辞或纠正错别字。所有全文目录与网页缓存均只保存在本地并通过.gitignore排除。公开仓库只保存来源元数据、派生统计和分析结果。"
        "本综合报告不列出180条来源URL，也不展示完整网页作文。",
    )
    add_callout(
        document,
        "语料命名",
        "作文网把文章归入高中作文栏目，但栏目归类不能独立证明作者身份、写作年龄或是否经过编辑。本报告使用“公开网络母语参照语料”，避免把来源标签误写成实验事实。",
        fill=LIGHT_FILL,
    )
    comparability = pd.DataFrame(
        [
            ["题目", "四个固定HSK题目", "按四题建立对应主题域", "部分母语组使用扩展或宽泛主题"],
            ["体裁", "J组记叙、Y组议论", "对应组优先匹配记叙或议论", "网站体裁标签可能不完全可靠"],
            ["篇幅", "约358至391字均值", "按学习者分位点匹配", "母语文本仍整体偏长"],
            ["写作条件", "语料库既有评分作文", "公开网页文章", "时间、修改和资料使用均不可比"],
            ["身份信息", "国籍可用", "栏目和年级标签可用", "网页作者身份不能独立核验"],
            ["文本处理", "标注清洗后分词", "网页正文清理后分词", "均使用相同301列统计口径"],
        ],
        columns=["维度", "学习者作文", "母语参照作文", "剩余限制"],
    )
    add_dataframe_table(
        document,
        comparability,
        columns=["维度", "学习者作文", "母语参照作文", "剩余限制"],
        widths_dxa=[1200, 2450, 2450, 3260],
        font_size=7.8,
    )
    add_table_source(document, "表2-1  两类语料的可比条件与不可消除差异。")

    document.add_page_break()
    document.add_heading("3 301项语言特征与统计方法", level=1)
    document.add_heading("3.1 特征体系", level=2)
    add_paragraph(
        document,
        "每篇作文的宽表共有301列，包括8列基本信息、5列基础篇幅指标和288列派生语言特征。派生特征覆盖词汇丰富度、词汇密度与词长、句段结构、词性、"
        "语法标记、篇章连接、记叙描写和HSK词汇八类。次数型特征同时保存原始次数与每千汉字频率，统一分母为纯文本汉字数；比例、均值、TTR、Guiraud和MATTR直接保存计算值。",
    )
    feature_system = pd.DataFrame(
        [
            ["基础篇幅", 5, "字数、纯文本字数、分词数、非标点分词数、去重词数"],
            ["词汇丰富度", 18, "TTR、Guiraud、MATTR-50、仅出现一次词及词类TTR"],
            ["词汇密度与词长", 21, "内容词、实虚词、词汇密度、平均词长及单双多字词"],
            ["句段结构", 22, "句段数、句长、长句、逗号、分句和段落长度"],
            ["词性", 46, "23类词性的次数与每千字频率"],
            ["语法标记", 56, "代词、情态、否定、体标记、结构助词、介词和语气词"],
            ["篇章连接", 21, "八类连接词及总量、去重数和多样性"],
            ["记叙描写", 26, "时间、动作、心理、言说、趋向、评价、引语和动词链"],
            ["HSK", 78, "数字等级、初中高、覆盖率、词种、TTR及非HSK拆分"],
        ],
        columns=["类别", "字段数", "核心内容"],
    )
    add_dataframe_table(
        document,
        feature_system,
        columns=["类别", "字段数", "核心内容"],
        widths_dxa=[1900, 1000, 6460],
        font_size=8.2,
    )
    add_table_source(document, "表3  301列宽表中的统计类别；另有8列基本信息。")
    add_paragraph(
        document,
        "HSK词汇使用11,000条新版考试大纲词表，主等级1至3归为初等、4至6归为中等、7至9归为高等。词汇占比分母为全部非标点token。"
        "同形词多等级时优先使用PyNLPIR词性匹配，仍无法唯一判断时归入最低主等级。非HSK词拆分为专名、数字、字母串和其他；“其他”不能直接解释为错误或高级词。",
    )

    document.add_heading("3.2 学习者五维模型", level=2)
    add_paragraph(
        document,
        "因子候选只使用每千字频率、比例、TTR、MATTR、Guiraud、均值和句长等可直接比较指标，排除原始次数、篇幅指标和确定性重复总量。"
        "依次剔除零方差、95%以上为零的稀疏变量、绝对相关不低于0.85的冗余变量及变量级KMO低于0.50的指标。候选模型使用MinRes提取和Promax斜交旋转，"
        "再以共同度、因子主要载荷数、Heywood异常和Bootstrap Tucker一致性筛选。",
    )
    screening_summary = pd.DataFrame(
        [
            ["初始直接指标", learner_meta["screening"]["initial_candidates"]],
            ["方差与稀疏性后", learner_meta["screening"]["after_variance_sparsity"]],
            ["相关性筛选后", learner_meta["screening"]["after_correlation"]],
            ["变量级KMO筛选后", learner_meta["screening"]["after_msa"]],
            ["最终五维模型", learner_meta["selected_model"]["feature_count"]],
        ],
        columns=["筛选阶段", "保留指标数"],
    )
    add_dataframe_table(
        document,
        screening_summary,
        columns=["筛选阶段", "保留指标数"],
        widths_dxa=[6000, 3360],
        font_size=9,
    )
    add_table_source(document, "表4  学习者MF/MD指标筛选流程。完整逐项原因见分析结果工作簿。")
    add_figure(document, learner_figures / "02_变量筛选流程.png", 3, "从301列宽表到最终39项指标的筛选流程")
    add_figure(document, learner_figures / "03_平行分析与碎石图.png", 4, "平行分析、碎石图与候选模型比较")
    add_paragraph(
        document,
        f"平行分析建议{learner_meta['parallel_analysis']['suggested_factor_count']}个因子，但高维方案无法同时满足可解释性、每因子主要载荷数和稳定性要求。"
        f"最终选定5因子、39变量模型，KMO={learner_meta['selected_model']['kmo']:.3f}，样本变量比为"
        f"{learner_meta['selected_model']['n_per_variable']:.1f}，200次Bootstrap的五个因子Tucker一致性中位数均高于"
        f"{min(learner_meta['selected_model']['stability_medians']):.3f}。",
    )

    document.add_heading("3.3 推断统计与固定投影", level=2)
    add_bullets(
        document,
        [
            "学习者四组比较使用Welch方差分析、Holm校正两两Welch检验、Omega平方和Hedges g。",
            "分数关系同时使用Spearman相关与HC3稳健回归：维度得分 ~ 作文分数 + 篇名代码 + 是否日本籍。",
            "母语参照使用学习者均值、标准差及载荷方向进行固定投影；学习者重新投影误差不超过4.44×10⁻¹⁶。",
            "四组母语对照报告Welch检验、Hedges g、Bootstrap区间和Holm校正；39项具体特征使用BH校正。",
            "稳健性包括1000次篇幅最近邻重抽样、严格主题与优先年份子样本、控制log篇幅的HC3回归和800篇联合EFA。",
        ],
    )
    metric_formulas = pd.DataFrame(
        [
            ["每千汉字频率", "特征次数 ÷ 纯文本字数 × 1000", "控制文本长度的机械影响"],
            ["TTR", "去重词数 ÷ 非标点分词数", "受文本长度影响的词形多样性"],
            ["Guiraud", "去重词数 ÷ √非标点分词数", "对篇幅进行平方根修正"],
            ["MATTR-50", "连续50词窗口TTR的平均值", "局部词汇丰富度；短文使用全文TTR"],
            ["HSK等级占比", "该等级token数 ÷ 非标点分词数", "词汇等级构成而非能力分数"],
            ["维度得分", "正载荷指标z分之和－负载荷指标z分之和", "再在学习者样本内标准化"],
        ],
        columns=["指标", "公式", "解释"],
    )
    add_dataframe_table(
        document,
        metric_formulas,
        columns=["指标", "公式", "解释"],
        widths_dxa=[1800, 3300, 4260],
        font_size=7.8,
    )
    add_table_source(document, "表4-1  核心指标及维度得分口径。")

    document.add_page_break()
    document.add_heading("4 学习者作文的五维结构", level=1)
    add_figure(document, learner_figures / "04_因子载荷热图.png", 5, "最终39项指标的Promax旋转载荷")
    add_paragraph(
        document,
        "载荷阈值取绝对值0.30。维度名称根据正负两端主要载荷及语言功能中性命名；斜交旋转允许维度相关，因此同一指标可能在多个维度上有次要载荷。"
        "下列解释使用主要载荷，而完整载荷矩阵见附录A。",
    )
    for index, dimension in enumerate(DIMENSIONS, start=1):
        document.add_heading(f"4.{index} D{index} {dimension}", level=2)
        local = assignments.loc[assignments["所属维度"] == dimension].copy()
        positive = local.loc[local["主要载荷"] > 0].sort_values("主要载荷", ascending=False).head(6)
        negative = local.loc[local["主要载荷"] < 0].sort_values("主要载荷").head(6)
        add_paragraph(document, DIMENSION_INTERPRETATIONS[dimension])
        positive_text = "、".join(f"{row['字段名']}（{row['主要载荷']:.2f}）" for _, row in positive.iterrows())
        negative_text = "、".join(f"{row['字段名']}（{row['主要载荷']:.2f}）" for _, row in negative.iterrows())
        add_paragraph(document, f"主要正载荷：{positive_text}。", bold_lead="主要正载荷：")
        if negative_text:
            add_paragraph(document, f"主要负载荷：{negative_text}。", bold_lead="主要负载荷：")
        cross = loadings.loc[loadings["主要维度"] == dimension].copy()
        add_dataframe_table(
            document,
            cross.sort_values("主要载荷", key=lambda values: values.abs(), ascending=False).head(8),
            columns=["字段名", "类别", "主要载荷", "共同度"],
            widths_dxa=[3300, 2200, 1700, 2160],
            font_size=8.2,
        )
        add_table_source(document, f"表4-{index}  {dimension}的主要归属指标。")

    model_comparison = model_diagnostics.loc[model_diagnostics["因子数"].between(2, 6)].copy()
    add_dataframe_table(
        document,
        model_comparison,
        columns=["因子数", "最终变量数", "KMO", "累计解释方差", "统计诊断通过", "稳定性通过", "诊断说明"],
        widths_dxa=[650, 950, 750, 1250, 1100, 1100, 3560],
        font_size=7.4,
    )
    add_table_source(document, "表4-6  2至6因子候选模型比较；五因子是通过诊断与稳定性检验的最高维度解。")
    add_paragraph(
        document,
        "五因子方案不是由平行分析单独决定，而是平衡统计诊断、Bootstrap稳定性、每因子至少三个主要载荷和可解释性后的结果。"
        "更高维模型虽可提高解释方差，却出现载荷支撑不足或稳定性下降；更低维模型则会把词汇、人称、句法和动作过程压缩到过宽的维度中。",
    )

    document.add_page_break()
    document.add_heading("5 学习者四组、分数与国籍", level=1)
    document.add_heading("5.1 四组多维画像", level=2)
    add_figure(document, learner_figures / "05_四组维度得分分布.png", 6, "J1、J2、Y1、Y2的五维得分分布")
    add_figure(document, learner_figures / "06_四组多维画像热图.png", 7, "四组作文的五维均值画像")
    learner_dim = learner_dimension_table(learner_descriptives, learner_welch)
    add_dataframe_table(
        document,
        learner_dim,
        columns=["维度", "J1", "J2", "Y1", "Y2", "Omega平方", "BH校正p值"],
        widths_dxa=[2600, 900, 900, 900, 900, 1200, 1960],
        font_size=8.0,
    )
    add_table_source(document, "表5  学习者四组五维均值及Welch总体效应。维度得分已在学习者样本内标准化。")
    add_paragraph(
        document,
        "五个维度的组间Welch检验均用于检验四组总体差异，显著结果还需结合两两比较和效应量判断。J1/J2与Y1/Y2的差异同时包含题目和体裁，"
        "因此报告使用“题目与体裁组合差异”而不使用纯体裁因果表述。",
    )
    add_figure(document, learner_figures / "09_组间效应量森林图.png", 10, "学习者四组两两比较的效应量")

    document.add_heading("5.2 作文分数", level=2)
    add_figure(document, learner_figures / "07_分数与维度得分趋势.png", 8, "作文分数与五维得分趋势")
    add_dataframe_table(
        document,
        score_correlations,
        columns=["维度名称", "Spearman_rho", "p值", "BH校正p值"],
        widths_dxa=[4200, 1700, 1700, 1760],
        font_size=8.5,
    )
    add_table_source(document, "表6  作文分数与五维得分的Spearman相关。")
    score_terms = learner_regressions.loc[learner_regressions["模型项"].astype(str).str.contains("作文分数")].copy()
    if not score_terms.empty:
        add_dataframe_table(
            document,
            score_terms,
            columns=["维度名称", "模型项", "系数", "稳健标准误", "95%CI下限", "95%CI上限", "BH校正p值"],
            widths_dxa=[2200, 1300, 950, 1100, 1100, 1100, 1610],
            font_size=7.8,
        )
        add_table_source(document, "表7  控制篇名代码和是否日本籍后的分数系数，标准误为HC3。")
    add_figure(document, learner_figures / "10_回归系数森林图.png", 11, "分数、题目组和日本籍控制变量的稳健回归系数")

    document.add_heading("5.3 HSK词汇与维度相关", level=2)
    add_figure(document, learner_figures / "08_HSK等级构成.png", 9, "学习者四组作文的HSK词汇等级构成")
    add_paragraph(
        document,
        "初等、中等、高等和非HSK词汇构成受题目词汇域显著影响。低等级词汇占比高并不等于语言简单：人物与旅行叙事需要大量基础动作、时间和空间词，"
        "而议论题更可能使用抽象名词和中高等级词汇。等级构成必须与TTR、词汇密度和五维功能共同解释。",
    )
    add_figure(document, learner_figures / "11_维度相关热图.png", 12, "五个斜交维度之间的相关结构")

    document.add_heading("5.4 日本籍与非日本籍补充分析", level=2)
    add_paragraph(
        document,
        "国籍补充分析仅比较日本籍与非日本籍，并控制题目和分数；不对42个国籍进行均值排名。由于日本籍在J2中的占比很高，国籍、题目和学习经历并非独立，"
        "回归中的国籍系数只能视为探索性控制项。任何显著差异都不能解释为民族或国籍本质属性。",
    )
    learner_internal_records: list[dict[str, Any]] = []
    for dimension in DIMENSIONS:
        local = learner_descriptives.loc[learner_descriptives["维度名称"] == dimension]
        high = local.loc[local["均值"].idxmax()]
        low = local.loc[local["均值"].idxmin()]
        welch_row = learner_welch.loc[learner_welch["维度名称"] == dimension].iloc[0]
        score_row = score_correlations.loc[score_correlations["维度名称"] == dimension].iloc[0]
        learner_internal_records.append(
            {
                "维度": dimension,
                "最高组": high["篇名代码"],
                "最低组": low["篇名代码"],
                "Omega平方": welch_row["Omega平方"],
                "分数rho": score_row["Spearman_rho"],
                "分数BH校正p值": score_row["BH校正p值"],
            }
        )
    learner_internal = pd.DataFrame(learner_internal_records)
    add_dataframe_table(
        document,
        learner_internal,
        columns=["维度", "最高组", "最低组", "Omega平方", "分数rho", "分数BH校正p值"],
        widths_dxa=[2800, 900, 900, 1400, 1300, 2060],
        font_size=8.0,
    )
    add_table_source(document, "表7-1  学习者内部证据摘要。最高/最低组按维度均值判定，不代表作文质量排序。")

    document.add_page_break()
    document.add_heading("6 公开网络母语参照样本审计", level=1)
    document.add_heading("6.1 来源、年级与年份", level=2)
    sources = source_summary(native_master)
    add_dataframe_table(
        document,
        sources,
        columns=["来源类型", "篇数", "占比"],
        widths_dxa=[5500, 1200, 2660],
        max_rows=10,
        font_size=8.3,
    )
    add_table_source(document, "表8  母语样本主表中的来源类型汇总；完整180行审计记录仍保留在Excel。")
    years = year_summary(native_master)
    add_dataframe_table(
        document,
        years,
        columns=["母语代码", "样本量", "最早年份", "最晚年份", "平均汉字数", "平均目标偏差", "年代扩展样本"],
        widths_dxa=[900, 800, 1000, 1000, 1400, 1900, 2360],
        font_size=8.0,
    )
    add_table_source(document, "表9  母语参照样本的发布年份、篇幅与年代扩展情况。发布日期不等同于实际写作年份。")
    add_paragraph(
        document,
        "网站来源标签包括作文网原创、网络资源、本站原创、转载等。标签只用于审计网页自述来源，不验证版权链或作者身份。"
        "采集脚本默认单线程、请求间隔3至6秒，并对429、403、5xx和超时进行指数退避；不绕过验证码或访问限制。",
    )

    document.add_heading("6.2 篇幅匹配", level=2)
    add_figure(document, native_figures / "02_篇幅匹配诊断.png", 13, "学习者与母语参照作文的篇幅分布及匹配诊断")
    length_summary = native_sampling[["母语代码", "平均汉字数", "中位汉字数", "平均目标篇幅相对偏差"]].copy()
    add_dataframe_table(
        document,
        length_summary,
        columns=["母语代码", "平均汉字数", "中位汉字数", "平均目标篇幅相对偏差"],
        widths_dxa=[1400, 1900, 1900, 4160],
        font_size=8.5,
    )
    add_table_source(document, "表10  母语参照样本篇幅概况。目标篇幅来自对应学习者组的分位点。")
    add_paragraph(
        document,
        "公开高中作文整体长于学习者作文。候选选择使用最小总偏差匹配，但网页文章的可用篇幅下界较高，无法实现完全平衡。"
        "这一差异可能机械抬高词种数、句法延展和部分篇章组织指标，因此后续同时使用每千字频率、log篇幅控制和篇幅匹配重抽样。",
    )

    document.add_heading("6.3 主题匹配", level=2)
    strict_counts = native_master["主题匹配层级"].value_counts().reindex(["精确", "近似", "扩展", "宽泛"], fill_value=0)
    add_paragraph(
        document,
        "180篇中，精确、近似、扩展和宽泛主题分别为"
        + "、".join(f"{name}{int(value)}篇" for name, value in strict_counts.items())
        + "。宽泛主题主要用于保持高中范围和四组等量，尤其影响NY1、NY2和NJ2。严格主题子样本用于方向敏感性检查，但样本量不足以独立承担显著性检验。",
    )
    add_callout(
        document,
        "母语主表在综合报告中的作用",
        "主表提供每篇网页文章的代码、URL、年级、发布日期、来源类型、主题层级、篇幅偏差和正文哈希。本报告只使用这些字段生成聚合审计信息，不把180行清单或网页全文复制进Word/PDF。",
    )
    sampling_risks = pd.DataFrame(
        [
            ["作者身份", "保留栏目、年级和来源元数据", "不能确认作者确为高中母语者"],
            ["主题偏差", "四级主题匹配并做严格主题敏感性", "Y1/Y2严格主题样本很少"],
            ["篇幅偏差", "分位点匹配、每千字频率、HC3与重抽样", "匹配后仍有残余长度差"],
            ["年代偏差", "优先2005至2012年并标记扩展样本", "发布日期不等同于写作年份"],
            ["重复文本", "SHA-256与字符五元组近重复去重", "不能识别所有人工改写版本"],
            ["编辑加工", "不修改正文并记录来源类型", "无法判断教师、编辑或平台修改程度"],
        ],
        columns=["风险", "已采取措施", "剩余不确定性"],
    )
    add_dataframe_table(
        document,
        sampling_risks,
        columns=["风险", "已采取措施", "剩余不确定性"],
        widths_dxa=[1500, 3800, 4060],
        font_size=7.8,
    )
    add_table_source(document, "表10-1  公开网络母语参照语料的主要风险与缓解措施。")

    document.add_page_break()
    document.add_heading("7 四组母语参照对照", level=1)
    document.add_heading("7.1 五维总体结果", level=2)
    add_figure(document, native_figures / "03_五维画像热图.png", 14, "学习者与母语参照八组作文的五维画像")
    add_figure(document, native_figures / "05_五维效应量森林图.png", 15, "四个题目中母语参照减学习者的五维Hedges g")
    comparison_display = native_comparisons.copy()
    add_dataframe_table(
        document,
        comparison_display,
        columns=[
            "对应题目",
            "维度",
            "学习者均值",
            "母语均值",
            "Hedges_g_母语减学习者",
            "Hedges_g_Bootstrap95%CI下限",
            "Hedges_g_Bootstrap95%CI上限",
            "Holm校正p值",
        ],
        widths_dxa=[800, 1850, 700, 700, 1100, 1100, 1200, 1910],
        font_size=6.9,
    )
    add_table_source(document, f"表11  四组×五维共20项比较；{native_significant}项经Holm校正后显著。正g表示母语参照更高。")

    for index, learner_code in enumerate(("J1", "J2", "Y1", "Y2"), start=2):
        document.add_heading(f"7.{index} {learner_code}/N{learner_code} {PAIR_LABELS[learner_code]}", level=2)
        local = native_comparisons.loc[native_comparisons["对应题目"] == learner_code].copy()
        local["绝对效应"] = local["Hedges_g_母语减学习者"].abs()
        local = local.sort_values("绝对效应", ascending=False)
        statements = []
        for _, row in local.iterrows():
            direction = "高于" if row["Hedges_g_母语减学习者"] > 0 else "低于"
            significance = "Holm显著" if row["Holm校正p值"] < 0.05 else "校正后不显著"
            statements.append(
                f"{row['维度']}{direction}学习者（g={row['Hedges_g_母语减学习者']:.2f}，{effect_label(row['Hedges_g_母语减学习者'])}效应，{significance}）"
            )
        add_paragraph(document, "；".join(statements) + "。")
        top_features = native_features.loc[native_features["对应题目"] == learner_code].copy()
        top_features["绝对效应"] = top_features["Hedges_g_母语减学习者"].abs()
        top_features = top_features.sort_values("绝对效应", ascending=False).head(6)
        add_paragraph(
            document,
            "具体指标中绝对效应较大的项目为："
            + "、".join(
                f"{row['字段名']}（g={row['Hedges_g_母语减学习者']:.2f}）" for _, row in top_features.iterrows()
            )
            + "。",
        )
        if learner_code == "J1":
            add_paragraph(document, "人物影响题应优先从人物指涉、评价资源和事件例证的组织方式解释。公开高中作文可能使用更多修辞和复句资源，但这也可能反映修改机会与网页筛选。")
        elif learner_code == "J2":
            add_paragraph(document, "假期与经历题要求时间、地点、趋向和动作序列。叙事推进及动作链是本组最具功能解释力的指标；学习者J2中的日本籍比例较高，限制了跨国籍外推。")
        elif learner_code == "Y1":
            add_paragraph(document, "Y1精确网页样本最少，大量参照文章来自健康行为与相邻社会议题。全样本差异必须与严格主题子样本方向一起阅读，不能视为严格同题结果。")
        else:
            add_paragraph(document, "Y2参照范围扩展至家风、责任、成长和教育方式。网页作文可能带有模板化议论结构，因此连接、词汇等级和句法差异都需结合来源条件解释。")

    document.add_heading("7.6 具体语言特征", level=2)
    add_figure(document, native_figures / "06_主要语言特征效应量.png", 16, "四组母语对照中绝对效应最大的具体语言特征")
    feature_summary = native_features.assign(
        绝对效应=lambda frame: frame["Hedges_g_母语减学习者"].abs()
    ).sort_values("绝对效应", ascending=False).head(20)
    add_dataframe_table(
        document,
        feature_summary,
        columns=["对应题目", "字段名", "所属维度", "学习者均值", "母语均值", "Hedges_g_母语减学习者", "BH校正p值"],
        widths_dxa=[700, 2100, 1800, 900, 900, 1200, 1760],
        font_size=7.2,
    )
    add_table_source(document, "表12  39项固定投影特征中绝对效应最大的20项；BH校正在156项比较内执行。")

    document.add_heading("7.7 HSK词汇构成", level=2)
    add_figure(document, native_figures / "07_HSK词汇构成.png", 17, "学习者与母语参照的HSK初中高及非HSK词汇构成")
    hsk = hsk_summary(learner_stats, native_stats)
    add_dataframe_table(
        document,
        hsk,
        columns=["代码", "初等词汇占比", "中等词汇占比", "高等词汇占比", "非HSK词汇占比", "HSK词汇覆盖率"],
        widths_dxa=[1000, 1400, 1400, 1400, 1700, 2460],
        font_size=7.8,
    )
    add_table_source(document, "表13  八组作文HSK词汇构成均值。所有比例的分母均为非标点分词数。")
    add_paragraph(
        document,
        "母语参照中的非HSK词汇包含专名、成语、新词、词表外表达和分词误差，不能直接标记为高级词。学习者与参照作文的HSK构成差异还受到题目词汇域影响，"
        "因此报告不使用单一覆盖率评价语言水平。",
    )
    hsk_interpretation = pd.DataFrame(
        [
            ["HSK覆盖率", "命中11,000条词表的token比例", "受专名、新词和分词影响"],
            ["高等词汇占比", "7至9级词汇占全部非标点token比例", "题目域可能机械抬高或压低"],
            ["等级内部TTR", "某等级词种数 ÷ 该等级token数", "低频等级样本过少时波动大"],
            ["非HSK词汇占比", "未命中词表的token比例", "不能直接解释为错误或高级词"],
            ["非HSK专名", "未命中词表且被识别为专名", "人物与地点题更容易出现"],
        ],
        columns=["指标", "回答问题", "解释限制"],
    )
    add_dataframe_table(
        document,
        hsk_interpretation,
        columns=["指标", "回答问题", "解释限制"],
        widths_dxa=[1900, 3700, 3760],
        font_size=8.0,
    )
    add_table_source(document, "表13-1  HSK派生指标的解释边界。")

    document.add_page_break()
    document.add_heading("8 稳健性与联合因子分析", level=1)
    document.add_heading("8.1 篇幅匹配重抽样", level=2)
    add_figure(document, native_figures / "08_篇幅匹配重抽样.png", 18, "1000次篇幅匹配重抽样的五维均值差")
    add_dataframe_table(
        document,
        resampling,
        columns=["对应题目", "维度", "重抽样均值差", "重抽样标准差", "篇幅均值差", "重抽样95%CI下限", "重抽样95%CI上限"],
        widths_dxa=[700, 2200, 1200, 1200, 1200, 1200, 1660],
        font_size=7.0,
    )
    add_table_source(document, f"表14  篇幅最近邻重抽样；{resampling_robust}/20项95%区间不跨0。")
    residuals = resampling.groupby("对应题目")["篇幅均值差"].first()
    add_paragraph(
        document,
        "匹配后的母语参照相对学习者平均残余篇幅差为"
        + "、".join(f"{code} {value:.1f}字" for code, value in residuals.items())
        + "。Y1残差最大，说明重抽样降低但没有消除篇幅混杂。",
    )

    document.add_heading("8.2 主题、年份与HC3回归", level=2)
    sensitivity_counts = (
        sensitivity.groupby("敏感性样本", as_index=False)
        .agg(比较项=("维度", "size"), 平均绝对差=("均值差_母语减学习者", lambda values: float(values.abs().mean())))
    )
    add_dataframe_table(
        document,
        sensitivity_counts,
        columns=["敏感性样本", "比较项", "平均绝对差"],
        widths_dxa=[4300, 1500, 3560],
        font_size=8.5,
    )
    add_table_source(document, "表15  全样本、严格主题和优先年份敏感性分析概况。")
    strict_n = (
        sensitivity.loc[sensitivity["敏感性样本"] == "精确或近似主题"]
        .groupby("对应题目")["母语样本量"]
        .first()
    )
    add_paragraph(
        document,
        "精确或近似主题子样本量为"
        + "、".join(f"{code}={int(value)}" for code, value in strict_n.items())
        + "。Y1和Y2样本量过小，只能用于方向检查。控制log篇幅的HC3模型进一步检验来源主效应及来源×题目交互，但仍无法控制作者背景、修改次数和编辑加工。",
    )
    source_terms = native_regressions.loc[native_regressions["项"].astype(str).str.contains("公开网络母语参照")].copy()
    source_terms["项"] = source_terms["项"].map(compact_native_regression_term)
    add_dataframe_table(
        document,
        source_terms,
        columns=["维度", "项", "系数", "HC3标准误", "95%CI下限", "95%CI上限", "BH校正p值"],
        widths_dxa=[1700, 2600, 800, 900, 900, 900, 1560],
        max_rows=20,
        font_size=7.0,
    )
    add_table_source(document, "表16  HC3模型中的语料来源主效应和题目交互项。")

    document.add_heading("8.3 联合因子结构", level=2)
    add_dataframe_table(
        document,
        joint_models,
        columns=["因子数", "变量数", "KMO", "累计解释方差", "诊断通过", "稳定性通过", "最小Tucker中位数", "诊断说明"],
        widths_dxa=[700, 700, 700, 1100, 900, 900, 1100, 3260],
        font_size=7.0,
    )
    add_table_source(document, "表17  800篇作文联合EFA候选模型诊断。")
    add_figure(document, native_figures / "10_联合模型Tucker一致性.png", 19, "学习者五维与800篇联合模型的Tucker一致性")
    joint = native_meta["joint_factor_sensitivity"]
    minimum_congruence = float(congruence["绝对Tucker一致性"].min()) if len(congruence) else 0.0
    add_paragraph(
        document,
        f"联合平行分析建议{joint['suggested_factors']}个因子，但2至10因子候选中只有{joint['selected_factors']}因子方案同时通过预设诊断和稳定性要求。"
        f"该方案保留{joint['selected_features']}项指标，KMO={joint['selected_kmo']:.3f}；与原五维的已匹配Tucker一致性最小值为{minimum_congruence:.2f}。"
        "因此联合结构不能完整复现原五维，固定五维投影适合描述差异，却不证明五维在两种语料来源中具有完全相同的潜在结构。",
    )
    projection_comparison = pd.DataFrame(
        [
            ["参数来源", "仅620篇学习者作文", "全部800篇作文"],
            ["主要目的", "保持原五维量尺并比较来源", "检查新增语料是否改变共现结构"],
            ["维度数量", "固定5维", f"通过诊断的方案为{joint['selected_factors']}维"],
            ["结果地位", "母语比较的主分析", "结构迁移的敏感性分析"],
            ["可解释结论", "同一量尺上的组间位置", "原五维能否跨来源完整重复"],
        ],
        columns=["项目", "固定五维投影", "联合EFA"],
    )
    add_dataframe_table(
        document,
        projection_comparison,
        columns=["项目", "固定五维投影", "联合EFA"],
        widths_dxa=[1700, 3830, 3830],
        font_size=8.0,
    )
    add_table_source(document, "表17-1  固定投影与联合因子分析承担的不同任务。")

    document.add_page_break()
    document.add_heading("9 五维综合证据", level=1)
    synthesis = synthesis_table(learner_descriptives, score_correlations, native_comparisons, resampling)
    add_dataframe_table(
        document,
        synthesis,
        columns=["维度", "学习者最高/最低", "分数rho", "最强母语对照", "Holm显著组数", "篇幅重抽样稳定组数"],
        widths_dxa=[2100, 1400, 1200, 2000, 1200, 1460],
        font_size=7.7,
    )
    add_table_source(document, "表18  学习者内部、分数关系和母语对照的五维证据汇总。")
    add_paragraph(
        document,
        "综合表把三类问题放在同一行：学习者组间位置、与分数的单调关系、母语参照的最大来源差异，以及篇幅重抽样后的稳定性。"
        "这些证据回答的是不同层级的问题，不能用一个显著结果替代另一个。例如，某维度与分数相关，并不意味着该维度上的母语参照方向就是评分标准。",
    )
    for index, dimension in enumerate(DIMENSIONS, start=1):
        document.add_heading(f"9.{index} {dimension}", level=2)
        learner = learner_descriptives.loc[learner_descriptives["维度名称"] == dimension].copy()
        high = learner.loc[learner["均值"].idxmax()]
        low = learner.loc[learner["均值"].idxmin()]
        score = score_correlations.loc[score_correlations["维度名称"] == dimension].iloc[0]
        native = native_comparisons.loc[native_comparisons["维度"] == dimension].copy()
        native["绝对效应"] = native["Hedges_g_母语减学习者"].abs()
        strongest = native.sort_values("绝对效应", ascending=False).iloc[0]
        significant_pairs = native.loc[native["Holm校正p值"] < 0.05, "对应题目"].tolist()
        add_paragraph(
            document,
            f"学习者内部均值最高组为{high['篇名代码']}（{high['均值']:.2f}），最低组为{low['篇名代码']}（{low['均值']:.2f}）。"
            f"与分数的Spearman相关为rho={score['Spearman_rho']:.2f}（BH校正后{p_text(score['BH校正p值'])}）。"
            f"母语参照最大差异出现在{strongest['对应题目']}，g={strongest['Hedges_g_母语减学习者']:.2f}，"
            f"Bootstrap 95%CI [{strongest['Hedges_g_Bootstrap95%CI下限']:.2f}, {strongest['Hedges_g_Bootstrap95%CI上限']:.2f}]。"
            + (f"Holm显著的题目为{'、'.join(significant_pairs)}。" if significant_pairs else "没有题目达到Holm校正后显著。"),
        )
        add_paragraph(document, DIMENSION_INTERPRETATIONS[dimension])

    document.add_heading("9.6 与参考论文的关系", level=2)
    literature = pd.DataFrame(
        [
            ["样本", "日本学习者与母语者叙述文", "多体裁学习者与母语者作文", "620篇HSK作文+180篇公开网页参照"],
            ["主要方法", "MF/MD因子分析", "MF/MD与多体裁比较", "学习者EFA、固定投影、稳健检验、联合EFA"],
            ["共同维度", "语句复杂性、动作描写", "词汇产出、动作描写等", "词汇丰富、句法延展、动作过程"],
            ["本项目特有", "单一叙述任务", "多体裁任务", "基础词汇推进、人称指涉及HSK等级"],
            ["关键限制", "学习者背景集中", "任务与体裁混合", "网页身份、主题扩展与篇幅不平衡"],
        ],
        columns=["项目", "徐勤2021", "徐勤2023博士论文", "本报告"],
    )
    add_dataframe_table(
        document,
        literature,
        columns=list(literature.columns),
        widths_dxa=[1100, 2600, 2700, 2960],
        font_size=7.7,
    )
    add_table_source(document, "表19  两篇参考论文与本综合报告的方法关系。")
    add_paragraph(
        document,
        "词汇丰富、句法延展和动作过程与参考研究有较清晰概念对应，基础词汇推进与人称指涉则更受本项目题目设计和HSK指标体系影响。"
        "这说明多维分析既能识别跨研究重复的功能成分，也会形成语料特定维度；维度命名必须依赖当前载荷而不能直接套用前人标签。",
    )
    evidence_priority = pd.DataFrame(
        [
            ["优先级1", "效应量区间、校正p值、重抽样和主题敏感性方向一致", "可进入人工文本复核"],
            ["优先级2", "主比较显著，但篇幅或严格主题结果不稳定", "保留结论并明确条件"],
            ["优先级3", "只有未经校正p值或单一描述指标支持", "作为探索线索，不写作主结论"],
            ["不作推断", "来源身份、编辑过程或写作条件无法核验", "不得转写为母语能力因果差异"],
        ],
        columns=["证据等级", "判定条件", "报告用途"],
    )
    add_dataframe_table(
        document,
        evidence_priority,
        columns=["证据等级", "判定条件", "报告用途"],
        widths_dxa=[1400, 5200, 2760],
        font_size=8.0,
    )
    add_table_source(document, "表19-1  综合结论的证据优先级。")

    document.add_page_break()
    document.add_heading("10 教学与研究启示", level=1)
    insights = [
        ("按题目功能组织教学。", "人物题需要指涉与评价，经历题需要时间、趋向和动作链，议论题需要抽象词汇、分句与逻辑关系。统一要求所有作文提高同一指标，容易偏离任务。"),
        ("把五维画像用于诊断。", "反馈可以分别指出词汇重复、人称过密、长句失控或动作链不足，而不把单一总分当作全部语言表现。维度高低应转化为可观察的语言选择。"),
        ("词汇等级与丰富度联合解释。", "HSK高等级占比、非HSK覆盖、MATTR和词汇密度分别回答不同问题。词汇等级高不保证篇内多样，非HSK多也不自动等于高级表达。"),
        ("优先复核多重证据一致的大效应。", "同时通过Holm/BH校正、Bootstrap、篇幅重抽样和主题敏感性的指标，适合作为后续人工文本分析与教学干预候选。"),
        ("建立同题、可核验的母语语料。", "下一轮应让身份和年级可核验的高中生在相同题目、时间限制和修改条件下写作，以区分语言背景、网页编辑与任务条件。"),
        ("保留模型的可复现链条。", "从清洗、分词、特征词表、301列宽表到39项因子模型均应版本化；新增语料时先固定投影，再把重新提取的因子结构作为敏感性分析。"),
    ]
    for lead, body in insights:
        add_paragraph(document, lead + body, bold_lead=lead)

    document.add_heading("11 局限", level=1)
    add_bullets(
        document,
        [
            "作文网栏目归类不能认证作者身份，也不能排除编辑、转载、润色或教师示例；来源差异不能直接解释为母语能力。",
            "NJ2、NY1和NY2为保持每组45篇而使用大量宽泛主题，尤其NY1严格主题样本极少；这些组更接近相邻题目域参照。",
            "母语参照整体长于学习者作文，最近邻匹配后仍存在残余篇幅差；每千字频率、HC3回归和重抽样只能部分缓解。",
            "两类语料的写作时间、修改机会、受众、是否使用资料和评分条件不同，来源变量混合了多种未观测情境。",
            "学习者作文没有75分层级，国籍样本量不均，日本籍与J2高度相关，限制了分数与国籍分析的外推。",
            "PyNLPIR会对专名、成语、新词和细粒度词性产生误差；共同处理口径支持相对比较，但不能保证每个token绝对正确。",
            "HSK词表中的非命中项不等于错误或高级词，同形多等级词的最低等级回退也会影响边界词统计。",
            "探索性因子分析依赖样本与指标集合；联合800篇语料只得到稳定的2因子方案，原五维不是跨来源完全不变的测量模型。",
            "统计显著不等于教学重要，效应量也不能替代具体文本功能；所有结论都需要人工语篇分析复核。",
            "网页全文因服务协议只在本地保存，公开仓库无法提供完全自足的文本复核包。",
        ],
    )

    document.add_heading("12 结论", level=1)
    add_paragraph(
        document,
        "本项目已经形成从620篇HSK学习者作文到180篇公开网络高中作文参照的统一分析链。301列宽表提供细粒度语言描述，39项最终指标把共同变异压缩为"
        "词汇丰富度与词汇扩展、基础词汇与叙事推进、人称指涉与信息密度、句法延展与分句复杂度、动作过程与动词链五个维度。学习者内部的题目组、分数和补充国籍分析"
        "与母语参照投影共同表明，作文语言差异是多方向的资源配置，而不是单一复杂度阶梯。",
    )
    add_paragraph(
        document,
        f"四组母语对照的20项维度比较中有{native_significant}项经Holm校正后显著，{resampling_robust}项在篇幅匹配重抽样中保持方向稳定。"
        "但宽泛主题、残余篇幅差和网页来源不确定性要求降低因果表述。联合因子分析未完整复制原五维，更进一步说明固定投影适合比较，不能替代独立验证。"
        "现阶段最合理的用途，是利用综合证据筛选值得人工复核和严格同题实验复现的语言特征。",
    )

    document.add_heading("参考文献与网页来源", level=1)
    references = [
        "徐勤（2021）．日本人中国語学習者の叙述文における言語的特徴の分析：MF/MD法を使って．《言語文化共同研究プロジェクト》2020，27-41．https://doi.org/10.18910/85007",
        "徐勤（2023）．中国語作文における言語特徴の考察：多次元分析による日本人中国語学習者と中国語母語話者の作文の比較．大阪大学博士学位论文．https://doi.org/10.18910/91825",
        "Biber, D. (1988). Variation Across Speech and Writing. Cambridge University Press.",
        "中华人民共和国教育部、国家语言文字工作委员会（2021）．《国际中文教育中文水平等级标准》（GF 0025—2021）。",
        "新版HSK考试大纲词汇表：项目文件《新版HSK考试大纲1219.pdf》及机器可读《新版HSK词汇大纲.csv》。",
        "作文网高中作文栏目：https://www.zuowen.com/gaozhong/。",
        "作文网服务协议：http://www.zuowen.com/help/agreement/。",
    ]
    for item in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(6)
        set_run_font(paragraph.add_run(item), size=9.8)

    document.add_page_break()
    document.add_heading("附录A 最终因子载荷", level=1)
    appendix_loadings = assignments.copy()
    appendix_loadings["维度序号"] = appendix_loadings["所属维度"].map(
        {name: index for index, name in enumerate(DIMENSIONS, start=1)}
    )
    appendix_loadings = appendix_loadings.sort_values(
        ["维度序号", "主要载荷"], key=lambda values: values.abs() if values.name == "主要载荷" else values,
        ascending=[True, False],
    )
    add_dataframe_table(
        document,
        appendix_loadings,
        columns=["字段名", "所属维度", "主要载荷", "共同度", "正负极"],
        widths_dxa=[2300, 2900, 1200, 1200, 1760],
        font_size=7.4,
    )
    add_table_source(document, "表A1  最终39项指标的主要维度归属。完整五维交叉载荷见《作文多维分析结果.xlsx》。")

    document.add_heading("附录B 母语样本审计摘要", level=1)
    audit = native_sampling.merge(
        years[["母语代码", "样本量", "最早年份", "最晚年份"]],
        on=["母语代码", "样本量"],
        how="left",
    )
    add_dataframe_table(
        document,
        audit,
        columns=[
            "母语代码",
            "样本量",
            "精确主题",
            "近似主题",
            "扩展主题",
            "宽泛主题",
            "最早年份",
            "最晚年份",
            "年代扩展样本",
        ],
        widths_dxa=[900, 700, 850, 850, 850, 850, 1000, 1000, 2360],
        font_size=7.8,
    )
    add_table_source(document, "表B1  四组母语参照样本的主题与年份审计。完整逐篇记录见《母语作文样本主表.xlsx》。")
    add_dataframe_table(
        document,
        sources,
        columns=["来源类型", "篇数", "占比"],
        widths_dxa=[5500, 1200, 2660],
        font_size=8.3,
    )
    add_table_source(document, "表B2  母语参照样本来源类型汇总。")

    document.add_heading("附录C 可复现性说明", level=1)
    reproducibility = pd.DataFrame(
        [
            ["学习者MF/MD随机种子", learner_meta["seed"]],
            ["学习者平行分析次数", learner_meta["parallel_analysis"]["iterations"]],
            ["学习者Bootstrap成功次数", learner_meta["selected_model"]["bootstrap_successes"]],
            ["母语分析随机种子", native_meta["random_seed"]],
            ["母语比较Bootstrap次数", native_meta["bootstrap_iterations"]],
            ["篇幅匹配重抽样次数", native_meta["length_resamples"]],
            ["联合平行分析次数", native_meta["parallel_iterations"]],
            ["联合Bootstrap次数", native_meta["joint_bootstrap"]],
        ],
        columns=["项目", "取值"],
    )
    add_dataframe_table(
        document,
        reproducibility,
        columns=["项目", "取值"],
        widths_dxa=[6500, 2860],
        font_size=8.5,
    )
    add_table_source(document, "表C1  两阶段分析的固定随机参数。")
    reproducibility_paragraph = add_paragraph(
        document,
        "复现顺序：先运行segment_hsk_clean_texts.py生成学习者宽表；运行analyze_composition_mfmd.py及其工作簿、报告脚本；"
        "再运行collect_zuowen_native_controls.py、segment_native_control_texts.py、analyze_native_control.py及其工作簿、报告脚本；"
        "最后运行build_integrated_composition_report.py。综合脚本只读取两阶段正式结果，不重估因子或改变样本。",
    )
    reproducibility_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    result_files_paragraph = add_paragraph(
        document,
        "主要结果文件包括《作文词性统计宽表.xlsx》《作文多维分析结果.xlsx》《母语作文样本主表.xlsx》《母语作文词性统计宽表.xlsx》"
        "和《作文母语对照分析结果.xlsx》。完整网页正文及缓存不进入Git，报告中的所有采样数字均可由母语样本主表重新汇总。",
    )
    result_files_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    build_report(
        Path(args.learner_analysis_dir).resolve(),
        Path(args.native_analysis_dir).resolve(),
        Path(args.learner_workbook).resolve(),
        Path(args.native_master).resolve(),
        Path(args.native_workbook).resolve(),
        Path(args.output).resolve(),
    )
    print(f"Report written: {Path(args.output).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

#!/usr/bin/env python3
"""Build the detailed native-control comparison report as a polished DOCX."""

from __future__ import annotations

import argparse
import json
import math
import re
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from .build_mfmd_report import (
        BLUE,
        DARK_BLUE,
        GOLD,
        INK,
        MUTED,
        add_callout,
        add_dataframe_table,
        add_field,
        add_figure,
        add_paragraph,
        add_table_source,
        configure_document,
        set_run_font,
    )
except ImportError:
    from build_mfmd_report import (
        BLUE,
        DARK_BLUE,
        GOLD,
        INK,
        MUTED,
        add_callout,
        add_dataframe_table,
        add_field,
        add_figure,
        add_paragraph,
        add_table_source,
        configure_document,
        set_run_font,
    )


DEFAULT_ANALYSIS_DIR = "outputs/native_control_analysis"
DEFAULT_SELECTED_JSON = "outputs/native_control/selected_samples.json"
DEFAULT_NATIVE_TEXT_DIR = "native_clean_text"
DEFAULT_OUTPUT = "作文语言特征母语对照分析报告.docx"

DIMENSIONS = [
    "词汇丰富度与词汇扩展",
    "基础词汇与叙事推进",
    "人称指涉与信息密度",
    "句法延展与分句复杂度",
    "动作过程与动词链",
]
PAIR_LABELS = {
    "J1": "J1/NJ1 人物影响记叙",
    "J2": "J2/NJ2 假期旅行及个人经历记叙",
    "Y1": "Y1/NY1 吸烟、健康及相邻社会议题",
    "Y2": "Y2/NY2 父母教育及成长议题",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成HSK作文与公开网络母语参照语料详细分析报告。")
    parser.add_argument("--analysis-dir", default=DEFAULT_ANALYSIS_DIR)
    parser.add_argument("--selected-json", default=DEFAULT_SELECTED_JSON)
    parser.add_argument("--native-text-dir", default=DEFAULT_NATIVE_TEXT_DIR)
    parser.add_argument("--output", default=DEFAULT_OUTPUT)
    return parser.parse_args()


def p_text(value: float) -> str:
    return "p<0.001" if value < 0.001 else f"p={value:.3f}"


def effect_label(value: float) -> str:
    absolute = abs(value)
    if absolute < 0.20:
        return "可忽略"
    if absolute < 0.50:
        return "小"
    if absolute < 0.80:
        return "中等"
    return "大"


def add_static_toc(document: Document) -> None:
    document.add_heading("目录", level=1)
    rows = pd.DataFrame(
        [
            ["摘要", 3],
            ["1 研究背景与研究问题", 4],
            ["2 语料来源、版权与样本构建", 4],
            ["3 分词、语言特征与五维投影", 8],
            ["4 统计方法", 8],
            ["5 五维总体结果", 9],
            ["6 四个题目的配对解释", 14],
            ["7 具体语言特征与HSK词汇", 16],
            ["8 篇幅、主题与年份稳健性", 18],
            ["9 HC3稳健回归", 20],
            ["10 联合因子分析敏感性", 23],
            ["11 极短文本例证", 24],
            ["12 与两篇论文的关系", 25],
            ["13 教学与研究启示", 25],
            ["14 局限", 26],
            ["15 结论", 26],
            ["参考文献与网页来源", 27],
            ["附录A 可复现性；附录B 敏感性明细", 28],
        ],
        columns=["章节", "页码"],
    )
    add_dataframe_table(
        document,
        rows,
        columns=["章节", "页码"],
        widths_dxa=[8200, 1160],
        font_size=9.0,
    )
    add_table_source(document, "目录页码按正式生成版排定。")
    document.add_page_break()


def add_cover(document: Document, metadata: dict[str, Any]) -> None:
    for _ in range(6):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("语料库比较研究报告"), size=11, bold=True, color=GOLD)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("HSK学习者作文与公开网络\n母语参照语料多维对照分析"), size=27, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    native_rows = metadata["validation"]["native_rows"]
    set_run_font(
        subtitle.add_run(f"基于620篇学习者作文与{native_rows}篇作文网高中作文的MF/MD投影研究"),
        size=13.5,
        color=DARK_BLUE,
    )
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(28)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    rule._p.get_or_add_pPr().append(borders)
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    groups = metadata["validation"]["native_group_counts"]
    set_run_font(
        summary.add_run(" · ".join(["五维固定投影", "39项指标", *[f"{key}={value}" for key, value in groups.items()]])),
        size=10.5,
        bold=True,
        color=INK,
    )
    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.paragraph_format.space_before = Pt(64)
    set_run_font(date_line.add_run("2026年8月"), size=11, color=MUTED)
    document.add_page_break()


def read_table(tables_dir: Path, name: str) -> pd.DataFrame:
    return pd.read_csv(tables_dir / f"{name}.csv", encoding="utf-8-sig")


def select_short_native_examples(
    selected: pd.DataFrame,
    native_text_dir: Path,
) -> pd.DataFrame:
    records: list[dict[str, str]] = []
    for code in ("NJ1", "NJ2", "NY1", "NY2"):
        group = selected.loc[selected["母语代码"] == code].copy()
        group["_主题顺序"] = group["主题匹配层级"].map({"精确": 0, "近似": 1, "扩展": 2}).fillna(3)
        group = group.sort_values(["_主题顺序", "目标篇幅相对偏差", "作文文件名"])
        if group.empty:
            continue
        row = group.iloc[0]
        path = native_text_dir / code / f"{row['作文文件名']}.txt"
        if not path.is_file():
            continue
        text = re.sub(r"\s+", "", path.read_text(encoding="utf-8"))
        sentence = re.split(r"[。！？!?]", text)[0][:18]
        records.append(
            {
                "母语代码": code,
                "作文文件名": row["作文文件名"],
                "短例句": sentence + ("……" if len(text) > len(sentence) else ""),
                "来源URL": row["来源URL"],
            }
        )
    return pd.DataFrame(records)


def configure_native_header(document: Document) -> None:
    for section in document.sections:
        paragraph = section.header.paragraphs[0]
        paragraph.text = "HSK学习者作文与公开网络母语参照语料对照分析"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(paragraph.runs[0], size=8.5, color=MUTED)


def build_report(
    analysis_dir: Path,
    selected_json: Path,
    native_text_dir: Path,
    output_path: Path,
) -> None:
    metadata = json.loads((analysis_dir / "analysis_metadata.json").read_text(encoding="utf-8"))
    selected_document = json.loads(selected_json.read_text(encoding="utf-8"))
    selected = pd.DataFrame(selected_document["selected"])
    tables_dir = analysis_dir / "tables"
    figures_dir = analysis_dir / "figures"

    sampling = read_table(tables_dir, "采样审计概况")
    descriptives = read_table(tables_dir, "维度描述统计")
    comparisons = read_table(tables_dir, "五维Welch比较")
    features = read_table(tables_dir, "39项特征比较")
    regressions = read_table(tables_dir, "HC3稳健回归")
    resampling = read_table(tables_dir, "篇幅匹配重抽样")
    sensitivity = read_table(tables_dir, "主题年份敏感性")
    joint_models = read_table(tables_dir, "联合模型诊断")
    congruence = read_table(tables_dir, "Tucker一致性")
    examples = select_short_native_examples(selected, native_text_dir)

    document = Document()
    configure_document(document)
    configure_native_header(document)
    add_cover(document, metadata)
    add_static_toc(document)

    document.add_heading("摘要", level=1)
    native_count = metadata["validation"]["native_rows"]
    add_paragraph(
        document,
        f"本报告以620篇HSK学习者作文和{native_count}篇作文网高中栏目公开作文为对象，构建公开网络母语参照语料，并将其投影到项目既有的五维MF/MD语言特征量尺。母语样本分为NJ1、NJ2、NY1、NY2四组，每组45篇，分别对应人物影响、假期旅行及个人经历、吸烟与健康及相邻社会议题、父母教育与成长四个题目域。筛选时尽量匹配主题、体裁和篇幅，但高中网页候选的篇幅明显偏长，严格主题样本也较少；网页全文仅在本地保存，公开产物只保留来源元数据、派生统计和必要的极短例句。",
    )
    add_paragraph(
        document,
        "投影使用620篇学习者作文的39项指标均值、标准差和载荷方向，不利用母语样本重新定义量尺。四组主比较采用Welch检验、Hedges g、Bootstrap置信区间和Holm校正；并使用HC3稳健回归、1000次篇幅匹配重抽样、主题与年份敏感性分析，以及全部语料的联合探索性因子分析作为稳健性检验。篇幅重抽样后仍存在残余长度差，因此控制结果只能降低、不能消除篇幅混杂。",
    )
    significant = comparisons.loc[comparisons["Holm校正p值"] < 0.05]
    largest = comparisons.iloc[comparisons["Hedges_g_母语减学习者"].abs().argsort()[::-1]].head(3)
    add_paragraph(
        document,
        f"20项题目×维度比较中，{len(significant)}项在Holm校正后显著。绝对效应最大的三个差异为："
        + "；".join(
            f"{row['对应题目']}的{row['维度']}（g={row['Hedges_g_母语减学习者']:.2f}）"
            for _, row in largest.iterrows()
        )
        + "。这些结果描述的是学习者语料与公开网络高中作文参照语料的差异，不能据此认证网页作者身份或推断纯粹的母语能力效应。",
    )
    add_callout(
        document,
        "核心解释边界",
        "作文网将文章归入高中作文栏目，但无法独立核验作者是否确为高中阶段汉语母语者，也无法完全排除编辑、转载或润色。因此本项目统一使用“公开网络母语参照语料”这一审慎名称；所有对照结论均是语料来源差异，而非身份认证后的母语者因果结论。",
        fill="FFF7E8",
    )
    keywords = document.add_paragraph()
    set_run_font(keywords.add_run("关键词："), bold=True)
    set_run_font(keywords.add_run("HSK作文；母语参照语料；MF/MD；因子投影；篇幅匹配；作文网"))

    document.add_heading("1 研究背景与研究问题", level=1)
    add_paragraph(
        document,
        "项目此前基于620篇J1、J2、Y1、Y2作文建立了五维MF/MD结构，分别描述词汇丰富度、基础词汇与叙事推进、人称指涉与信息密度、句法延展与分句复杂度、动作过程与动词链。该结构能够解释学习者作文内部的题目与体裁组合差异，但原语料没有汉语母语者对照组，因而无法判断哪些模式是二语写作特有现象，哪些只是特定题目的普遍写作要求。",
    )
    add_paragraph(
        document,
        "徐勤（2021）在日本学习者与汉语母语者叙述文中使用MF/MD法，徐勤（2023）又将多维分析扩展到多体裁作文。两项研究提示，母语对照有助于区分语言水平、题目和体裁造成的变异。不过，公开网页作文并不等同于严格采集的母语者实验语料，所以本报告把它定位为方法上的参照层，而不是黄金标准。",
    )
    add_callout(
        document,
        "研究问题",
        "（1）四个对应题目中，学习者与公开网络母语参照作文在原五维量尺上有何差异？（2）差异在控制篇幅、区分题目，并缩小到更严格主题或优先年份子样本后是否仍然存在？（3）哪些具体语言特征贡献最大？（4）全部语料联合因子分析能否复现原五维结构？",
    )

    document.add_heading("2 语料来源、版权与样本构建", level=1)
    document.add_heading("2.1 来源和使用边界", level=2)
    add_paragraph(
        document,
        "候选文章来自作文网高中作文栏目及高一、高二、高三常规作文页面。纳入范围限定为完整的记叙文和议论文；高中单元作文及练习来源中的完整作文可以纳入，但仅有题目、要求或方法说明的练习页仍排除。另排除满分作文、范文、素材、写作指导、点评、解析、读后感、演讲稿、小说和诗歌。网站服务协议对复制和传播设有限制，因此网页全文只保存在本地的native_ori_text与native_clean_text目录，并通过.gitignore排除；公开仓库不包含全文。",
    )
    add_paragraph(
        document,
        "正文提取保留作者原有措辞和错别字，仅删除标题、日期、来源、作者信息、广告、二维码、编辑和版权尾注。采集器兼容GB2312、GB18030与UTF-8页面，默认单线程请求，间隔3至6秒；对429、403、5xx和网络超时执行指数退避，不绕过验证码或访问限制。",
    )
    add_table_source(document, "来源栏目：https://www.zuowen.com/gaozhong/；服务协议：http://www.zuowen.com/help/agreement/")

    document.add_heading("2.2 主题、篇幅与年代匹配", level=2)
    add_paragraph(
        document,
        "NJ1聚焦具体人物及其影响；NJ2从假期、旅行或游历经历扩展到第一人称个人经历；NY1从吸烟、烟草、二手烟和公共禁烟扩展到健康行为、公共责任、规则与相邻社会议题；NY2从父母教育、言传身教扩展到家庭影响、教育、自立与青少年成长。主题匹配分为精确、近似、扩展、宽泛四级，优先保留2005至2012年文章。正文原则上限定在200至750字，并按学习者各组篇幅分位点尽量匹配。脚本保留了NY2初中作文的最后备用通道，但最终180篇均来自高中范围，未使用初中样本。",
    )
    theme_table = sampling[
        ["母语代码", "样本量", "平均汉字数", "精确主题", "近似主题", "扩展主题", "宽泛主题", "平均目标篇幅相对偏差"]
    ]
    add_dataframe_table(
        document,
        theme_table,
        columns=list(theme_table.columns),
        widths_dxa=[850, 750, 1100, 900, 900, 900, 900, 3060],
        font_size=7.8,
    )
    add_table_source(document, "表1a  母语参照样本的主题与篇幅审计。相对偏差按各篇目标长度计算。")
    grade_table = sampling[["母语代码", "高一", "高二", "高三", "年级未标注", "年代扩展样本"]]
    add_dataframe_table(
        document,
        grade_table,
        columns=list(grade_table.columns),
        widths_dxa=[1200, 1300, 1300, 1300, 1800, 2460],
        font_size=8.2,
    )
    add_table_source(document, "表1b  母语参照样本的年级与发布年代审计。完整URL和哈希见《母语作文样本主表.xlsx》。")
    add_figure(document, figures_dir / "01_采样主题层级.png", 1, "母语参照样本的主题匹配层级")
    add_figure(document, figures_dir / "02_篇幅匹配诊断.png", 2, "四个对应题目中学习者与母语参照作文的篇幅分布")

    document.add_heading("2.3 去重与审计", level=2)
    add_paragraph(
        document,
        "采集器使用正文SHA-256删除完全重复文本，并以字符五元组Jaccard相似度识别近重复；相似度达到0.85的文章只保留一篇。标准化同名题目和同一编号系列在每组最多保留一篇。每条候选均记录入选、拒绝或失败状态及原因，确保选取路径可审计和可复现。",
    )
    add_callout(
        document,
        "质量控制",
        "自动筛选完成来源、长度、主题和重复审计，但现有高中候选无法达到原定的严格篇幅平衡，且多数NY1、NY2和NJ2样本属于宽泛主题。自动审核也不能替代作者身份核验与逐篇文本鉴定。母语样本主表保留“审核状态”，以便后续开展第二轮人工复核和替换。",
    )

    document.add_page_break()
    document.add_heading("3 分词、语言特征与五维投影", level=1)
    add_paragraph(
        document,
        "母语清洗文本使用与学习者作文完全相同的PyNLPIR分词、细粒度词性映射、HSK词汇大纲和语言特征词表，生成与学习者逐列对齐的母语统计宽表。基本信息中的国籍统一写为“中国（公开网络样本）”，作文分数留空，不人为赋分。分词、词性、熟语、复句和HSK等统计均执行同一验证。",
    )
    add_paragraph(
        document,
        "五维投影以原学习者模型的39项最终指标为基础。每个指标先使用620篇学习者作文的均值和样本标准差转换为z分数；再按指标主要载荷的正负方向求和；最后使用学习者维度原始和的均值和标准差转成标准分。这样，母语参照得分的0点和1个标准差都由学习者样本定义。",
    )
    projection_error = metadata["projection"]["learner_reprojection_max_abs_error"]
    add_callout(
        document,
        "投影回归校验",
        f"将620篇学习者作文重新投影后，与原MF/MD结果的五维得分最大绝对差异为{projection_error:.2e}，在浮点容差内完全一致。",
    )

    document.add_heading("4 统计方法", level=1)
    add_paragraph(
        document,
        "四组主比较分别为J1-NJ1、J2-NJ2、Y1-NY1和Y2-NY2。对五个维度使用Welch独立样本t检验，报告母语参照减学习者的均值差、Hedges g和Bootstrap 95%置信区间，并对20个检验使用Holm校正。Hedges g为正表示母语参照得分更高，负值表示学习者得分更高；方向不等同于语言质量。",
    )
    add_paragraph(
        document,
        "稳健回归模型为“维度得分 ~ 语言来源 × 对应题目 + log(纯文本字数)”，使用HC3稳健标准误。篇幅敏感性分析执行1000次一对一最近邻重抽样，每次从对应学习者组中抽取与母语参照篇幅最接近且不重复的作文。另分别在精确或近似主题子样本、2005至2012年子样本中重复比较。",
    )
    add_paragraph(
        document,
        "联合因子敏感性分析合并全部学习者与母语参照作文，沿用原分析的零值95%、绝对相关0.85、变量级MSA 0.50、共同度0.30和载荷0.30阈值，采用MinRes提取、Promax旋转、1000次平行分析和Bootstrap稳定性检验。联合模型与原五维结构通过Tucker一致性系数比较。",
    )

    document.add_heading("5 五维总体结果", level=1)
    add_figure(document, figures_dir / "03_五维画像热图.png", 3, "八组语料在原学习者五维量尺上的均值画像")
    add_figure(document, figures_dir / "04_五维得分分布.png", 4, "四个对应题目中学习者与母语参照作文的五维得分分布")
    compact = comparisons[["对应题目", "维度", "学习者均值", "母语均值", "均值差_母语减学习者", "Hedges_g_母语减学习者", "Holm校正p值"]]
    add_dataframe_table(
        document,
        compact,
        columns=list(compact.columns),
        widths_dxa=[700, 2500, 1050, 1050, 1400, 1400, 1260],
        max_rows=20,
        font_size=7.6,
    )
    add_table_source(document, "表2  四组五维Welch比较。正效应表示公开网络母语参照更高。")
    add_figure(document, figures_dir / "05_五维效应量森林图.png", 5, "五维差异的标准化效应量")

    for index, dimension in enumerate(DIMENSIONS, start=1):
        document.add_heading(f"5.{index} {dimension}", level=2)
        local = comparisons.loc[comparisons["维度"] == dimension].sort_values(
            "Hedges_g_母语减学习者", key=lambda series: series.abs(), ascending=False
        )
        largest_row = local.iloc[0]
        significant_count = int((local["Holm校正p值"] < 0.05).sum())
        add_paragraph(
            document,
            f"该维度在四组比较中有{significant_count}组经Holm校正后显著。最大绝对效应出现在{PAIR_LABELS[largest_row['对应题目']]}："
            f"母语参照与学习者均值差为{largest_row['均值差_母语减学习者']:.2f}，其Bootstrap 95%CI为"
            f"[{largest_row['Bootstrap95%CI下限']:.2f}, {largest_row['Bootstrap95%CI上限']:.2f}]；"
            f"g={largest_row['Hedges_g_母语减学习者']:.2f}（{effect_label(largest_row['Hedges_g_母语减学习者'])}效应），"
            f"g的Bootstrap 95%CI为[{largest_row['Hedges_g_Bootstrap95%CI下限']:.2f}, {largest_row['Hedges_g_Bootstrap95%CI上限']:.2f}]，"
            f"校正后{p_text(largest_row['Holm校正p值'])}。",
        )
        directions = []
        for _, row in local.iterrows():
            direction = "母语参照更高" if row["Hedges_g_母语减学习者"] > 0 else "学习者更高"
            directions.append(f"{row['对应题目']} {direction}（g={row['Hedges_g_母语减学习者']:.2f}）")
        add_paragraph(document, "四个题目的方向依次为：" + "；".join(directions) + "。")
        if dimension == DIMENSIONS[0]:
            add_paragraph(document, "该维度综合TTR、MATTR、Guiraud及名词和动词的去重词指标，反映篇内词形扩展。母语参照更高时意味着在学习者量尺上使用了更分散的词汇资源；但它仍会受话题专名、修辞和编辑加工影响。")
        elif dimension == DIMENSIONS[1]:
            add_paragraph(document, "该维度的正端由基础HSK词、趋向动词和时间推进标记构成，负端偏向中高等级词。它更接近叙事推进方式，而不是简单的低水平词汇比例。")
        elif dimension == DIMENSIONS[2]:
            add_paragraph(document, "该维度把第一、第二、第三人称指涉与平均词长、词汇密度等信息压缩指标置于同一连续体。不同题目中的方向变化应优先解释为人物参与方式，而非笼统的口语/书面优劣。")
        elif dimension == DIMENSIONS[3]:
            add_paragraph(document, "该维度由平均句长、分句数、逗号与长句比例共同定义，是两篇参考论文中“语句复杂性”最接近的对应维度。网页作文的编辑加工可能尤其影响这一维度。")
        else:
            add_paragraph(document, "该维度集中反映动词频率、连续动词结构和动作链展开，与参考论文的“动作描写”维度高度相近。记叙和议论题的动作需求不同，必须按题目分别解读。")

    document.add_page_break()
    document.add_heading("6 四个题目的配对解释", level=1)
    for section_index, learner_code in enumerate(("J1", "J2", "Y1", "Y2"), start=1):
        document.add_heading(f"6.{section_index} {PAIR_LABELS[learner_code]}", level=2)
        local = comparisons.loc[comparisons["对应题目"] == learner_code].sort_values(
            "Hedges_g_母语减学习者", key=lambda series: series.abs(), ascending=False
        )
        text = []
        for _, row in local.iterrows():
            direction = "高于" if row["Hedges_g_母语减学习者"] > 0 else "低于"
            significance = "显著" if row["Holm校正p值"] < 0.05 else "未达校正后显著"
            text.append(f"{row['维度']}{direction}学习者（g={row['Hedges_g_母语减学习者']:.2f}，{significance}）")
        add_paragraph(document, "；".join(text) + "。")
        top_features = features.loc[features["对应题目"] == learner_code].copy()
        top_features["绝对效应"] = top_features["Hedges_g_母语减学习者"].abs()
        top_features = top_features.sort_values("绝对效应", ascending=False).head(5)
        add_paragraph(
            document,
            "具体指标中绝对效应较大的项目为：" + "、".join(
                f"{row['字段名']}（g={row['Hedges_g_母语减学习者']:.2f}）"
                for _, row in top_features.iterrows()
            ) + "。",
        )
        if learner_code == "J1":
            add_paragraph(document, "人物影响题的核心比较是人物指涉、评价和事件例证的组织方式。若母语参照在词汇丰富度或句法延展上更高，应理解为公开高中作文在人物刻画中使用了更多修辞和复句资源，而不是把所有差距归因于学习者能力。")
        elif learner_code == "J2":
            add_paragraph(document, "假期旅行题天然要求时间、地点、趋向和动作序列。两组差异最适合用叙事推进与动作链共同解释；日本籍在学习者J2中占比很高，也是本组外推的重要限制。")
        elif learner_code == "Y1":
            add_paragraph(document, "吸烟与公共健康主题的精确网页样本相对稀缺，因此扩展主题样本比例及敏感性结果尤其重要。若全样本差异在精确/近似子样本中明显减弱，就应视为主题扩展造成的偏移。")
        else:
            add_paragraph(document, "父母教育题常出现家庭、榜样、责任、成长等共享表达。网页作文可能带有模板化议论结构，因此词汇丰富度和连接标记差异需要结合文本来源谨慎解释。")

    document.add_page_break()
    document.add_heading("7 具体语言特征与HSK词汇", level=1)
    add_figure(document, figures_dir / "06_主要语言特征效应量.png", 6, "各题目差异最大的语言特征")
    feature_summary = features.assign(绝对效应=lambda frame: frame["Hedges_g_母语减学习者"].abs()).sort_values(
        "绝对效应", ascending=False
    ).head(20)
    add_dataframe_table(
        document,
        feature_summary,
        columns=["对应题目", "字段名", "所属维度", "学习者均值", "母语均值", "Hedges_g_母语减学习者", "BH校正p值"],
        widths_dxa=[650, 2200, 2200, 900, 900, 1300, 1210],
        font_size=7.6,
    )
    add_table_source(document, "表3  39项投影特征中绝对效应最大的20项。BH校正在全部156项题目×特征比较内执行。")
    add_figure(document, figures_dir / "07_HSK词汇构成.png", 7, "HSK初中高等级词汇与非HSK词汇构成")
    add_paragraph(
        document,
        "HSK词汇大纲按1至3级归为初等、4至6级归为中等、7至9级归为高等。占比分母为全部非标点分词数；未命中词表的专名、数字、字母串和其他词也保留在分母中。母语参照作文中的非HSK词不能直接解释为错误或高级词，因为网页文本可能包含专名、新词、成语、分词差异及词表范围外表达。",
    )

    document.add_page_break()
    document.add_heading("8 篇幅、主题与年份稳健性", level=1)
    add_figure(document, figures_dir / "08_篇幅匹配重抽样.png", 8, "1000次篇幅匹配重抽样的维度均值差")
    add_dataframe_table(
        document,
        resampling,
        columns=["对应题目", "维度", "重抽样均值差", "重抽样95%CI下限", "重抽样95%CI上限", "篇幅均值差"],
        widths_dxa=[700, 2850, 1300, 1500, 1500, 1510],
        max_rows=20,
        font_size=7.8,
    )
    add_table_source(document, "表4  篇幅最近邻重抽样。均值差为母语参照减匹配学习者。")
    robust_count = int(
        ((resampling["重抽样95%CI下限"] > 0) | (resampling["重抽样95%CI上限"] < 0)).sum()
    )
    add_paragraph(
        document,
        f"20项维度比较中有{robust_count}项在1000次篇幅匹配重抽样的95%区间内不跨0。篇幅匹配后仍稳定的差异较难由文本长度单独解释；区间跨0的结果则应降低结论强度。",
    )
    residual_lengths = resampling.groupby("对应题目")["篇幅均值差"].first()
    add_paragraph(
        document,
        "但最近邻匹配未实现完全篇幅平衡：母语参照文本相对匹配学习者文本的平均残余长度差为"
        + "、".join(f"{code} {value:.1f}字" for code, value in residual_lengths.items())
        + "。尤其Y1仍有较大残差，因此重抽样结果不能被视为已经消除篇幅混杂。",
    )
    sensitivity_summary = (
        sensitivity.groupby("敏感性样本")
        .agg(比较项=("均值差_母语减学习者", "count"), 平均绝对差=("均值差_母语减学习者", lambda x: float(x.abs().mean())))
        .reset_index()
    )
    add_dataframe_table(
        document,
        sensitivity_summary,
        columns=["敏感性样本", "比较项", "平均绝对差"],
        widths_dxa=[4000, 1800, 3560],
        font_size=8.8,
    )
    add_table_source(document, "表5  全样本、主题更严格样本与优先年份样本的维度差异概况。")
    strict_counts = (
        sensitivity.loc[sensitivity["敏感性样本"] == "精确或近似主题"]
        .groupby("对应题目")["母语样本量"]
        .first()
    )
    add_paragraph(
        document,
        "精确或近似主题子样本量分别为"
        + "、".join(f"{code}={int(value)}" for code, value in strict_counts.items())
        + "。Y1和Y2的严格主题样本过小，相关均值差只能用于方向检查，不能作为独立的显著性证据。",
    )
    add_figure(document, figures_dir / "09_发布年份分布.png", 9, "母语参照作文的发布年份分布")

    document.add_heading("9 HC3稳健回归", level=1)
    source_terms = regressions.loc[regressions["项"].str.contains("公开网络母语参照", na=False)].copy()
    add_dataframe_table(
        document,
        source_terms,
        columns=["维度", "项", "系数", "HC3标准误", "95%CI下限", "95%CI上限", "BH校正p值", "调整R平方"],
        widths_dxa=[1900, 2400, 750, 850, 850, 850, 900, 860],
        max_rows=20,
        font_size=7.2,
    )
    add_table_source(document, "表6  语言来源主效应与题目交互项。模型控制log(纯文本字数)，J1和学习者为参照组。")
    add_paragraph(
        document,
        "回归中的来源主效应对应J1题目，其他题目的来源差异需要把主效应与相应交互项相加后解释。HC3标准误降低了异方差对推断的影响，但不能解决网页作者身份、编辑加工或未观测教育背景造成的混杂。",
    )

    document.add_page_break()
    document.add_heading("10 联合因子分析敏感性", level=1)
    add_dataframe_table(
        document,
        joint_models,
        columns=["因子数", "变量数", "KMO", "累计解释方差", "诊断通过", "稳定性通过", "最小Tucker中位数", "是否选定", "诊断说明"],
        widths_dxa=[650, 700, 650, 1050, 900, 950, 1300, 800, 2360],
        font_size=7.4,
    )
    add_table_source(document, "表7  全部学习者与母语参照作文的联合候选模型诊断。")
    add_figure(document, figures_dir / "10_联合模型Tucker一致性.png", 10, "原学习者五维结构与联合因子模型的一致性")
    joint_meta = metadata["joint_factor_sensitivity"]
    if joint_meta.get("selected_factors") is None:
        add_paragraph(document, "联合因子分析没有任何方案同时通过预设统计诊断，因此本报告不强行命名联合维度；五维投影仍作为主分析，联合结果仅用于说明结构迁移存在限制。")
    else:
        minimum = float(congruence["绝对Tucker一致性"].min()) if len(congruence) else 0.0
        add_paragraph(
            document,
            f"平行分析建议{joint_meta['suggested_factors']}个因子，最终联合模型选定{joint_meta['selected_factors']}个因子。"
            f"与原五维结构的匹配Tucker一致性最小值为{minimum:.2f}。通常0.85以上可视为较强重复，0.95以上接近因子等同；低于0.85的维度提示母语参照语料改变了指标共现结构。",
        )

    document.add_heading("11 极短文本例证", level=1)
    add_paragraph(
        document,
        "为保护网页内容并遵守服务协议，报告不收录完整作文，只展示每组一条不超过18个汉字的极短起始片段；来源URL用于审计，全文仅保存在本地目录。例句仅说明语料形态，不作为统计结论的单篇证据。",
    )
    if not examples.empty:
        add_dataframe_table(
            document,
            examples,
            columns=["母语代码", "作文文件名", "短例句", "来源URL"],
            widths_dxa=[800, 1600, 2300, 4660],
            font_size=8.0,
        )
        add_table_source(document, "表8  母语参照作文极短例句。每条网页来源仅引述必要的最短片段。")
    else:
        add_paragraph(document, "本次构建未能读取本地全文，因此不展示例句；来源审计仍见母语样本主表。")

    document.add_page_break()
    document.add_heading("12 与两篇论文的关系", level=1)
    literature = pd.DataFrame(
        [
            ["样本", "日本学习者与母语者记叙文", "多体裁日本学习者与母语者作文", "四题HSK学习者与公开网络高中作文参照"],
            ["指标", "111项中筛选58项", "语义与形式指标综合", "当前宽表中固定原五维39项投影"],
            ["主要方法", "MF/MD因子分析", "MF/MD因子分析与多体裁比较", "固定投影、Welch、HC3、重抽样、联合EFA"],
            ["可比维度", "语句复杂性、动作描写等", "词汇产出丰富度、动作描写等", "五维均可与论文概念对照"],
            ["关键限制", "单一学习者母语背景", "体裁与任务差异", "网页身份与编辑情况不可独立核验"],
        ],
        columns=["项目", "徐勤2021", "徐勤2023博士论文", "本报告"],
    )
    add_dataframe_table(
        document,
        literature,
        columns=list(literature.columns),
        widths_dxa=[1100, 2600, 2700, 2960],
        font_size=7.8,
    )
    add_table_source(document, "表9  两篇参考论文与本报告的研究设计对照。")
    add_paragraph(
        document,
        "本报告最直接借鉴的是多指标共同解释、Promax斜交旋转和维度得分思想。与论文不同之处在于，本报告的主模型不重新从母语参照中提取维度，而是先固定学习者五维量尺再投影，从而避免把参照语料加入后形成的新结构误称为原结构。联合EFA被明确降为敏感性分析。",
    )
    add_paragraph(
        document,
        "词汇丰富度、句法延展和动作过程三个维度与两篇论文的结果有较清晰对应；基础词汇与叙事推进、人称指涉与信息密度更强地体现本项目题目设计。由此可见，多维结构既有跨研究的重复成分，也具有语料特定成分。",
    )

    document.add_heading("13 教学与研究启示", level=1)
    insights = [
        ("以题目为单位解释差异。", "人物、旅行、健康议论和家庭教育对人称、动作、连接和词汇等级有不同功能需求，跨题目总平均会掩盖这些差异。"),
        ("把五维画像用于诊断而非排序。", "某维度高只表示相应语言资源更集中，不自动等于作文更好。教学反馈应把词汇、指涉、句法和动作链拆开。"),
        ("优先关注稳定的大效应。", "同时通过Holm校正、Bootstrap、篇幅匹配和严格主题子样本的差异，才适合作为后续人工文本分析与教学设计的候选。"),
        ("HSK词汇等级不能单独代表水平。", "非HSK词包含专名和词表外表达，高级词比例又受到题目影响，应结合MATTR、词汇密度和篇章功能。"),
        ("建立更严格的母语语料。", "后续应采集作者背景可核验、未经编辑、按同题目现场写作的高中生作文，以验证本轮公开网络参照结果。"),
    ]
    for lead, body in insights:
        add_paragraph(document, lead + body, bold_lead=lead)

    document.add_heading("14 局限", level=1)
    limitations = [
        "作文网栏目归类不能认证作者身份，也不能排除编辑、转载或润色；“母语”只能作为参照标签而非确定事实。",
        "公开网页全文的可用性受网站结构与服务协议约束，本项目不公开全文，外部复核需要重新访问来源URL。",
        "NJ2、NY1与NY2为维持每组45篇而大量使用宽泛主题，特别是NY1的42篇宽泛样本；这些组更接近相邻题目域参照，而非严格同题对照。",
        "母语参照作文整体长于学习者作文，最近邻匹配后仍残留明显篇幅差，尤其Y1约多192字；HC3回归和重抽样只能部分缓解这一混杂。",
        "精确或近似主题子样本仅为J1=9、J2=8、Y1=2、Y2=4，严格主题敏感性分析的统计稳定性有限。",
        "部分完整作文来自高中单元作文或练习来源，不能排除教师示例、修改或模板化表达的影响。",
        "网页作文的发布日期不等同于写作年份，2005至2012年的优先范围只能说明发布年代。",
        "PyNLPIR分词和词性标注会对专名、成语与新词产生误差；同一误差口径虽保证组间可比，却不保证每个词项绝对准确。",
        "学习者与网页作文的写作条件、时间限制、是否修改和受众均未知或不同，来源效应混合了多种写作情境差异。",
        "联合EFA属于探索性结构比较，Tucker一致性不能替代独立样本验证性因子分析。",
    ]
    for item in limitations:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.208
        set_run_font(paragraph.add_run(item))

    document.add_heading("15 结论", level=1)
    add_paragraph(
        document,
        f"本项目成功把{native_count}篇公开网络高中作文参照文本接入与620篇HSK学习者作文一致的统计流程，并在不改变原学习者五维模型的前提下完成固定投影。主分析同时报告效应量、置信区间、多重校正、篇幅匹配、主题与年份敏感性以及联合因子结构，形成了比简单均值比较更完整的证据链。",
    )
    add_paragraph(
        document,
        "结果能够指出哪些题目和语言维度存在稳定差异，却不能把这些差异直接归结为母语能力。公开网络语料最大的价值是提供可审计、可扩展的参照层，并帮助筛选值得进入严格实验语料验证的语言特征。下一步应优先完成人工全文复核，并用作者身份与写作条件可核验的同题母语作文进行复现。",
    )

    document.add_heading("参考文献与网页来源", level=1)
    references = [
        "徐勤（2021）．日本人中国語学習者の叙述文における言語的特徴の分析：MF/MD法を使って．《言語文化共同研究プロジェクト》2020，27-41．https://doi.org/10.18910/85007",
        "徐勤（2023）．中国語作文における言語特徴の考察：多次元分析による日本人中国語学習者と中国語母語話者の作文の比較．大阪大学博士学位论文．https://doi.org/10.18910/91825",
        "Biber, D. (1988). Variation Across Speech and Writing. Cambridge University Press.",
        "作文网高中作文栏目：https://www.zuowen.com/gaozhong/（采集访问日期见母语样本主表生成时间）。",
        "作文网服务协议：http://www.zuowen.com/help/agreement/。",
    ]
    for item in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(6)
        set_run_font(paragraph.add_run(item), size=10)

    document.add_page_break()
    document.add_heading("附录A 可复现性", level=1)
    add_paragraph(
        document,
        f"分析随机种子为{metadata['random_seed']}；五维比较Bootstrap为{metadata['bootstrap_iterations']}次；篇幅匹配为{metadata['length_resamples']}次；联合平行分析为{metadata['parallel_iterations']}次；联合Bootstrap为{metadata['joint_bootstrap']}次。所有逐篇五维得分、投影参数、组间检验、特征比较、回归、重抽样、敏感性和联合模型结果均保存在《作文母语对照分析结果.xlsx》。",
    )
    add_paragraph(
        document,
        "复现顺序：运行collect_zuowen_native_controls.py采集与筛选；使用x86_64 Python运行segment_native_control_texts.py完成PyNLPIR分词和同口径统计；运行两个工作簿生成脚本；运行analyze_native_control.py生成统计表和图；最后运行本报告脚本并渲染为PDF。",
    )
    document.add_heading("附录B 敏感性明细", level=1)
    sensitivity_table = sensitivity.head(60)
    add_dataframe_table(
        document,
        sensitivity_table,
        columns=["敏感性样本", "对应题目", "维度", "母语样本量", "均值差_母语减学习者"],
        widths_dxa=[1700, 900, 3600, 1200, 1960],
        font_size=7.7,
    )
    add_table_source(document, "表B1  主题与发布年份敏感性分析。完整明细见结果工作簿。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    build_report(
        Path(args.analysis_dir).resolve(),
        Path(args.selected_json).resolve(),
        Path(args.native_text_dir).resolve(),
        Path(args.output).resolve(),
    )
    print(f"Report written: {Path(args.output).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

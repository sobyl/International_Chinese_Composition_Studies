#!/usr/bin/env python3
"""Build the Chinese MF/MD analysis report from reproducible analysis outputs."""

from __future__ import annotations

import argparse
import json
import math
from datetime import date
from pathlib import Path
from typing import Any, Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DEFAULT_ANALYSIS_DIR = "outputs/mfmd_analysis"
DEFAULT_OUTPUT = "作文语言特征多维分析报告.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667580"
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
RULE = "B8C6D0"
WHITE = "FFFFFF"
GOLD = "A67C2D"
RED = "9B1C1C"
DOC_FONT = "Arial Unicode MS"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成HSK作文MF/MD中文分析报告。")
    parser.add_argument("--analysis-dir", default=DEFAULT_ANALYSIS_DIR)
    parser.add_argument("--output", default=DEFAULT_OUTPUT)
    return parser.parse_args()


def set_run_font(
    run: Any,
    *,
    ascii_name: str = DOC_FONT,
    east_asia_name: str = DOC_FONT,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = ascii_name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ascii_name)
    fonts.set(qn("w:hAnsi"), ascii_name)
    fonts.set(qn("w:eastAsia"), east_asia_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in CELL_MARGIN_DXA.items():
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell: Any, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"表格列宽之和必须为{CONTENT_WIDTH_DXA} DXA：{widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def prevent_row_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    properties.append(marker)


def add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = DOC_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in [style.name for style in document.styles]:
        caption = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = document.styles["Figure Caption"]
    caption.font.name = DOC_FONT
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    if "Source Note" not in [style.name for style in document.styles]:
        source = document.styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = document.styles["Source Note"]
    source.font.name = DOC_FONT
    source.font.size = Pt(8.5)
    source.font.color.rgb = RGBColor.from_string(MUTED)
    source._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    source._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    source._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)

    for current_section in document.sections:
        header = current_section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.text = "HSK作文语言特征 MF/MD 多维分析"
        set_run_font(paragraph.runs[0], size=8.5, color=MUTED)
        footer = current_section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer_paragraph.add_run("第 ")
        set_run_font(run, size=8.5, color=MUTED)
        add_field(footer_paragraph, "PAGE")
        run = footer_paragraph.add_run(" 页")
        set_run_font(run, size=8.5, color=MUTED)
    set_update_fields(document)


def add_paragraph(
    document: Document,
    text: str,
    *,
    bold_lead: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    before: float = 0,
    after: float = 8,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.333
    paragraph.alignment = alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run_font(first, bold=True)
        second = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(second)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_FILL) -> None:
    table = document.add_table(rows=2, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_cell_shading(table.cell(0, 0), DARK_BLUE)
    title_paragraph = table.cell(0, 0).paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(0)
    run = title_paragraph.add_run(title)
    set_run_font(run, size=10, bold=True, color=WHITE)
    set_cell_shading(table.cell(1, 0), fill)
    body = table.cell(1, 0).paragraphs[0]
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = 1.25
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = body.add_run(text)
    set_run_font(run, size=10)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def display_value(value: Any, header: str = "") -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (int,)):
        return f"{value:,}"
    if isinstance(value, float):
        if "p值" in header:
            return "<0.001" if value < 0.001 else f"{value:.3f}"
        if any(marker in header for marker in ("占比", "比例", "覆盖率", "方差")):
            return f"{value:.1%}"
        return f"{value:.3f}"
    return str(value)


def add_dataframe_table(
    document: Document,
    frame: pd.DataFrame,
    *,
    columns: Sequence[str],
    widths_dxa: Sequence[int],
    max_rows: int | None = None,
    font_size: float = 8.5,
) -> Any:
    selected = frame.loc[:, list(columns)]
    if max_rows is not None:
        selected = selected.head(max_rows)
    table = document.add_table(rows=1, cols=len(columns))
    set_table_geometry(table, widths_dxa)
    set_repeat_header(table.rows[0])
    for index, header in enumerate(columns):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=INK)
    for _, row in selected.iterrows():
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, header in enumerate(columns):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            if isinstance(row[header], (int, float)) and not isinstance(row[header], bool):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(display_value(row[header], header))
            set_run_font(run, size=font_size)
    return table


def add_table_source(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Source Note")
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, color=MUTED)


def add_figure(document: Document, path: Path, number: int, caption: str, width: float = 6.25) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", caption)
    caption_paragraph = document.add_paragraph(style="Figure Caption")
    caption_run = caption_paragraph.add_run(f"图{number}  {caption}")
    set_run_font(caption_run, size=9, color=MUTED)


def add_cover(document: Document, metadata: dict[str, Any]) -> None:
    for _ in range(5):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("语料库多维统计研究报告")
    set_run_font(run, size=11, bold=True, color=GOLD)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("HSK作文语言特征\nMF/MD多维分析")
    set_run_font(run, size=28, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("基于J1、J2、Y1、Y2四类620篇作文的探索性研究")
    set_run_font(run, size=14, color=DARK_BLUE)
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(30)
    properties = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    properties.append(borders)
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_after = Pt(8)
    run = summary.add_run(
        f"最终模型：{metadata['selected_model']['factor_count']}个语言维度 · "
        f"{metadata['selected_model']['feature_count']}项指标 · "
        f"KMO={metadata['selected_model']['kmo']:.3f}"
    )
    set_run_font(run, size=11, bold=True, color=INK)
    report_date = document.add_paragraph()
    report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_date.paragraph_format.space_before = Pt(70)
    run = report_date.add_run("2026年8月")
    set_run_font(run, size=11, color=MUTED)
    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("目录", level=1)
    toc_rows = pd.DataFrame(
        [
            ["摘要", 3],
            ["1 研究背景与研究问题", 3],
            ["2 数据来源与样本结构", 4],
            ["3 语言指标与统计方法", 6],
            ["4 MF/MD五维结构", 8],
            ["5 四个题目组的多维语言画像", 15],
            ["6 作文分数与语言维度", 18],
            ["7 日本籍与非日本籍的补充分析", 20],
            ["8 HSK等级词汇构成", 22],
            ["9 与两篇论文的对照", 23],
            ["10 教学与研究启示", 24],
            ["11 局限与后续工作", 25],
            ["12 结论", 26],
            ["参考文献", 27],
            ["附录", 28],
        ],
        columns=["章节", "页码"],
    )
    add_dataframe_table(
        document,
        toc_rows,
        columns=["章节", "页码"],
        widths_dxa=[8400, 960],
        font_size=9.5,
    )
    note = document.add_paragraph(style="Source Note")
    run = note.add_run("目录页码对应本次正式生成版本。")
    set_run_font(run, size=8.5, color=MUTED)
    document.add_page_break()


def top_loading_text(detail: dict[str, Any]) -> str:
    positive = "、".join(
        f"{item['字段名']}（{item['载荷']:.2f}）" for item in detail["正向主要指标"][:4]
    )
    negative = "、".join(
        f"{item['字段名']}（{item['载荷']:.2f}）" for item in detail["负向主要指标"][:4]
    )
    if not negative:
        negative = "未出现绝对载荷达到0.30的稳定负向指标"
    return f"正向载荷主要包括{positive}；负向载荷主要包括{negative}。"


def significance_text(p_value: float) -> str:
    if p_value < 0.001:
        return "p<0.001"
    return f"p={p_value:.3f}"


def build_report(analysis_dir: Path, output_path: Path) -> None:
    metadata = json.loads((analysis_dir / "analysis_metadata.json").read_text(encoding="utf-8"))
    tables_dir = analysis_dir / "tables"

    def read_table(name: str) -> pd.DataFrame:
        return pd.read_csv(tables_dir / f"{name}.csv")

    sample_summary = read_table("样本概况")
    screening = read_table("变量筛选")
    diagnostics = read_table("候选模型诊断")
    stability = read_table("因子稳定性")
    loadings = read_table("因子载荷")
    assignment = read_table("指标维度归属")
    score_sensitivity = read_table("得分敏感性")
    dimension_descriptive = read_table("维度描述统计")
    omnibus = read_table("Welch检验")
    pairwise = read_table("两两比较")
    score_correlations = read_table("分数相关")
    regressions = read_table("稳健回归")
    examples = read_table("匿名文本例证")
    dimension_scores = read_table("维度得分")

    selected = metadata["selected_model"]
    column_count = metadata.get("validation", {}).get("columns", "当前")
    labels = selected["dimension_labels"]
    details = selected["dimension_details"]

    document = Document()
    configure_document(document)
    add_cover(document, metadata)
    add_toc(document)

    document.add_heading("摘要", level=1)
    add_paragraph(
        document,
        f"本报告基于HSK作文语料库中J1、J2、Y1、Y2四类作文各155篇、共620篇文本，使用项目既有的{column_count}列语言特征宽表开展多特征/多维度（MF/MD）探索性分析。研究以每千汉字频率、词汇丰富度、句段结构、复句关系、叙事描写及HSK等级词汇指标为基础，经稀疏性、共线性、变量级MSA和共同度筛选后，采用MinRes提取与Promax斜交旋转形成最终多维结构。",
    )
    add_paragraph(
        document,
        f"最终保留{selected['feature_count']}项语言指标并提取{selected['factor_count']}个维度，整体KMO为{selected['kmo']:.3f}，Bartlett球形检验显著（χ²={selected['bartlett_chi2']:.1f}，df={selected['bartlett_df']}，p<0.001），累计解释方差为{selected['cumulative_variance']:.1%}。五个维度分别概括为：{ '；'.join(f'D{i + 1} {label}' for i, label in enumerate(labels)) }。Bootstrap结果显示各维度载荷结构具有很高的重复稳定性。",
    )
    add_paragraph(
        document,
        "四组作文在五个维度上的差异均达到统计显著，其中人称指涉与信息密度、基础词汇与叙事推进两个维度的组间效应最大。作文分数与多数维度的简单相关较弱；控制篇名代码和日本籍后，分数对句法延展维度呈小幅正向关联，对基础词汇与叙事推进维度呈小幅负向关联。日本籍与非日本籍之间的补充差异必须结合题目分布不均衡谨慎解释。",
    )
    add_callout(
        document,
        "核心结论",
        "本语料的主要变异首先由题目与体裁组合驱动，而不是由作文分数单独驱动。五维结构揭示了不同写作任务在词汇资源、人称组织、句法展开和动作表达上的系统差异；这些维度适合用作后续作文分类、评分解释和教学诊断指标，但不应被直接视为单一的“好/坏作文”尺度。",
    )
    keywords = document.add_paragraph()
    run = keywords.add_run("关键词：")
    set_run_font(run, bold=True)
    run = keywords.add_run("HSK作文；MF/MD；探索性因子分析；词汇丰富度；句法复杂度；二语写作")
    set_run_font(run)

    document.add_heading("1 研究背景与研究问题", level=1)
    add_paragraph(
        document,
        "MF/MD方法通过因子分析识别语言指标之间的共现模式，再将统计因子解释为语言变异维度。徐勤（2021）以日本汉语学习者与汉语母语者记叙文为对象，从111项语言特征中筛选58项并提取7个维度；徐勤（2023）的博士论文进一步扩展到约1450篇、五类体裁作文，最终提出词汇产出丰富度、状态情境表达、副词修饰和动作描写等4个维度。两项研究共同说明，单个词频指标难以完整刻画学习者写作，维度化的语言画像更适合解释复杂的共现关系。",
    )
    add_paragraph(
        document,
        "本项目当前语料与上述研究存在两点关键差异。第一，当前样本全部来自HSK作文语料，没有汉语母语者对照组；第二，J1/J2为记叙文题目，Y1/Y2为议论文题目，题目和体裁并非完全可分离。因此，本报告借鉴论文的统计流程与解释框架，但不复刻“日本学习者与母语者”的结论，也不把J/Y差异解释为纯粹的体裁因果效应。",
    )
    add_callout(
        document,
        "研究问题",
        "（1）620篇作文中哪些语言指标形成稳定的共现维度？（2）J1、J2、Y1、Y2四类作文的多维语言画像有何差异？（3）作文分数与各维度之间是否存在稳定关联？（4）控制题目和分数后，日本籍与非日本籍作文是否仍表现出维度差异？",
    )

    document.add_heading("2 数据来源与样本结构", level=1)
    add_paragraph(
        document,
        f"分析数据来自项目根目录的《作文词性统计宽表.xlsx》。该表由清洗后的620篇作文通过PyNLPIR分词及项目语言特征脚本生成，每篇作文对应一行，包含基本信息、篇幅、词汇丰富度、词汇密度与词长、句段结构、词性、熟语、语法标记、复句关系、记叙描写和HSK等级等{column_count}列。次数型指标均同时提供每千汉字频率，降低文本长度差异造成的机械影响。",
    )
    add_dataframe_table(
        document,
        sample_summary,
        columns=["篇名代码", "篇名", "体裁", "样本量", "平均分", "平均纯文本字数", "国籍数"],
        widths_dxa=[700, 3600, 900, 700, 750, 1700, 1010],
        font_size=8.3,
    )
    add_table_source(document, "表1  样本结构。四组分数分层完全一致，每组均为155篇。")
    add_figure(document, analysis_dir / "figures/01_样本与分数结构.png", 1, "样本量与分数分层结构")
    add_paragraph(
        document,
        "抽样采用按分数段对齐的设计，因此四个篇名代码在55、60、65、70、80、85、90和95分上的篇数完全一致。75分样本已从当前主表中排除，报告中的分数趋势因此不能用于判断70分到80分之间的连续变化。该平衡设计有利于比较题目组，但各分数段样本量并不均衡，55分仅4篇、70分192篇。",
    )

    document.add_page_break()
    document.add_heading("3 语言指标与统计方法", level=1)
    document.add_heading("3.1 指标筛选", level=2)
    add_paragraph(
        document,
        f"因子分析从{metadata['screening']['initial_candidates']}项候选指标开始。候选范围包括每千汉字频率以及TTR、Guiraud、MATTR、词汇密度、比例、均值、中位数、标准差和最长序列等直接指标；原始次数、篇幅变量及可由其他列确定性求和得到的总量不进入因子模型。随后依次剔除零方差变量、零值比例达到95%的稀疏变量、绝对Pearson相关系数达到0.85的冗余变量和变量级MSA低于0.50的指标。",
    )
    add_figure(document, analysis_dir / "figures/02_变量筛选流程.png", 2, "MF/MD候选变量筛选流程")
    screening_counts = screening.groupby(["筛选状态", "筛选原因"]).size().reset_index(name="变量数")
    add_dataframe_table(
        document,
        screening_counts,
        columns=["筛选状态", "筛选原因", "变量数"],
        widths_dxa=[1300, 6600, 1460],
        max_rows=12,
        font_size=8.5,
    )
    add_table_source(document, "表2  变量筛选状态与主要删除原因。完整逐字段记录见分析结果工作簿。")

    document.add_heading("3.2 因子提取、旋转与模型选择", level=2)
    add_paragraph(
        document,
        "本报告使用最小残差法（MinRes）提取公因子，并使用Promax斜交旋转，允许语言维度彼此相关。因子数量同时参考Horn平行分析、碎石图与候选模型诊断。对2至10因子方案分别迭代剔除共同度低于0.30的变量；合格模型须满足总体KMO不低于0.60、Bartlett检验显著、样本变量比不低于5、每个因子至少3个绝对载荷达到0.30的主要指标且不存在Heywood异常。",
    )
    add_figure(document, analysis_dir / "figures/03_平行分析与碎石图.png", 3, "碎石图、平行分析与最终维度数")
    add_dataframe_table(
        document,
        diagnostics,
        columns=["因子数", "最终变量数", "KMO", "累计解释方差", "每因子主要指标数", "Heywood异常", "统计诊断通过", "稳定性通过", "是否选定"],
        widths_dxa=[650, 900, 700, 1050, 2100, 900, 1050, 950, 1060],
        font_size=7.9,
    )
    add_table_source(
        document,
        "表3  2至10因子候选模型诊断。6因子起出现Heywood异常，因此选取通过全部门槛的最高维度解，即5维模型。",
    )
    add_paragraph(
        document,
        f"平行分析提示的统计上限为{metadata['parallel_analysis']['suggested_factor_count']}个因子，但高维方案并不自动等于可解释方案。6至10因子解出现共同度超过1或独特性为负的Heywood异常，8至10因子解还存在主要指标少于3个的维度。最终5因子解保留{selected['feature_count']}项指标，样本变量比为{selected['n_per_variable']:.2f}，各因子Bootstrap Tucker一致性中位数均高于{min(selected['stability_medians']):.3f}。",
    )

    document.add_heading("3.3 维度得分与推断统计", level=2)
    add_paragraph(
        document,
        "主要维度得分采用徐勤（2023）的可解释口径：将最终指标转为z分数，对正载荷指标求和、对负载荷指标反向求和，再将维度总分标准化。作为敏感性检验，同时计算因子分析器的回归型因子得分。四组比较使用Welch方差分析和Holm校正的两两Welch检验，并报告Omega平方与Hedges g；分数关系使用Spearman相关及带HC3稳健标准误的回归模型。多重检验按分析族使用Benjamini-Hochberg或Holm校正。",
    )

    document.add_page_break()
    document.add_heading("4 MF/MD五维结构", level=1)
    add_figure(document, analysis_dir / "figures/04_因子载荷热图.png", 4, "最终39项指标的Promax因子载荷矩阵", width=6.15)
    add_figure(document, analysis_dir / "figures/11_维度相关热图.png", 5, "Promax斜交旋转后的维度相关")
    add_paragraph(
        document,
        "五个维度不是彼此完全独立的能力分量，而是语言资源在真实作文中的共现方向。载荷正负表示维度两端，不能直接解释为“高分/低分”或“正确/错误”。下文对每个维度分别结合主要载荷、题目组均值、评分关系和代表性文本进行解释。",
    )

    for factor, (label, detail) in enumerate(zip(labels, details, strict=True), start=1):
        document.add_heading(f"4.{factor} D{factor} {label}", level=2)
        add_paragraph(document, top_loading_text(detail))
        mean_rows = dimension_descriptive[dimension_descriptive["维度名称"] == label]
        mean_map = dict(zip(mean_rows["篇名代码"], mean_rows["均值"], strict=True))
        ordered = sorted(mean_map, key=mean_map.get, reverse=True)
        omnibus_row = omnibus.loc[omnibus["维度名称"] == label].iloc[0]
        score_row = score_correlations.loc[score_correlations["维度名称"] == label].iloc[0]
        add_paragraph(
            document,
            f"组均值从高到低依次为{'、'.join(f'{code}（{mean_map[code]:.2f}）' for code in ordered)}。"
            f"Welch检验显示四组存在差异，F({omnibus_row['分子自由度']:.0f}, {omnibus_row['分母自由度']:.1f})={omnibus_row['Welch_F']:.2f}，"
            f"{significance_text(float(omnibus_row['BH校正p值']))}，Omega平方={omnibus_row['Omega平方']:.3f}。"
            f"作文分数的Spearman相关为ρ={score_row['Spearman_rho']:.3f}，校正后{significance_text(float(score_row['BH校正p值']))}。",
        )
        factor_examples = examples[examples["维度名称"] == label].sort_values("维度位置")
        for _, example in factor_examples.iterrows():
            add_callout(
                document,
                f"{example['维度位置']}例证 · {example['作文文件名']} · {int(example['作文分数'])}分 · 维度得分{example['维度得分']:.2f}",
                str(example["作文片段"]),
                fill="FAFBFC",
            )
        if factor == 1:
            add_paragraph(
                document,
                "该维度集中反映同一篇作文中不同词形的扩展程度。TTR、MATTR、Guiraud、名词/动词TTR和仅出现一次词占比共同提高，意味着文本在控制篇幅后使用了更分散的词汇资源。它与徐勤（2023）提出的“产出性词汇丰富度”最为接近。",
            )
        elif factor == 2:
            add_paragraph(
                document,
                "该维度的正端同时出现1至2级词汇、趋向动词和时间词，负端则聚集中等及高等HSK词汇，因而更适合解释为基础词汇支撑的事件推进，而不是简单的词汇水平高低。J2在该维度显著较高，符合假期叙事对时间推进和趋向表达的需求。",
            )
        elif factor == 3:
            add_paragraph(
                document,
                "该维度正端突出第一/第三人称代词和代词总体频率，负端则对应平均词长、非HSK词汇及词汇密度。它呈现“人物参与和指涉”与“压缩的信息密度”之间的连续对照，J1的人物叙述题在正端最突出。",
            )
        elif factor == 4:
            add_paragraph(
                document,
                "该维度由平均每句分句数、平均句长、句长离散程度、长句和逗号频率构成，直接反映句子内部的延展和分句组织。它与徐勤（2021）的“语句复杂性”维度高度接近，也是本报告中与作文分数呈稳定小幅正向联系的维度。",
            )
        else:
            add_paragraph(
                document,
                "该维度以动词频率、连续动词结构和情态动词为核心，反映事件与行动链的密集展开。它与两篇论文中的“动作描写”维度形成明确对应，但高得分只表示动作过程表达更集中，并不自动意味着整体作文质量更高。",
            )

    document.add_page_break()
    document.add_heading("5 四个题目组的多维语言画像", level=1)
    add_figure(document, analysis_dir / "figures/05_四组维度得分分布.png", 6, "四组作文在五个维度上的得分分布")
    add_figure(document, analysis_dir / "figures/06_四组多维画像热图.png", 7, "J1、J2、Y1、Y2的五维均值画像")
    group_means = dimension_descriptive.pivot(index="维度名称", columns="篇名代码", values="均值").reset_index()
    add_dataframe_table(
        document,
        group_means,
        columns=["维度名称", "J1", "J2", "Y1", "Y2"],
        widths_dxa=[3760, 1400, 1400, 1400, 1400],
        font_size=8.7,
    )
    add_table_source(document, "表4  五个标准化维度在四个题目组中的均值。全体样本均值为0。")
    add_paragraph(
        document,
        "J1的鲜明特征是人称指涉与信息密度维度最高，同时词汇丰富度也高于总体平均。这与“对我影响最大的一个人”的人物中心叙述相吻合。J2在基础词汇与叙事推进维度最高，但动作过程维度最低，说明假期题更多依赖时间、趋向和基础叙事资源组织经历，而不一定形成密集的连续动作链。",
    )
    add_paragraph(
        document,
        "Y1在动作过程与动词链维度最高、人称指涉维度最低，显示吸烟影响题的论述较少围绕具体人物，却大量使用表示行为、作用和可能性的动词。Y2在词汇丰富度维度最低，其余维度更接近总体平均；这可能意味着父母教育题存在较强的共享论述模板和高频表达，但需要结合具体词项和例文进一步验证。",
    )
    add_figure(document, analysis_dir / "figures/09_组间效应量森林图.png", 8, "四组差异较大的维度效应量")

    document.add_page_break()
    document.add_heading("6 作文分数与语言维度", level=1)
    add_figure(document, analysis_dir / "figures/07_分数与维度得分趋势.png", 9, "各分数段与五个维度的组内均值趋势")
    score_regressions = regressions[regressions["模型项"] == "score_z"].copy()
    add_dataframe_table(
        document,
        score_regressions,
        columns=["维度名称", "系数", "稳健标准误", "95%CI下限", "95%CI上限", "BH校正p值", "调整R平方"],
        widths_dxa=[3100, 900, 1050, 1050, 1050, 1100, 1110],
        font_size=8.2,
    )
    add_table_source(document, "表5  控制篇名代码和日本籍后的作文分数系数。作文分数已标准化。")
    significant_scores = score_regressions[score_regressions["BH校正p值"] < 0.05]
    if significant_scores.empty:
        score_text = "控制题目和日本籍后，没有维度与作文分数保持多重校正后的显著联系。"
    else:
        score_text = "控制题目和日本籍后，" + "；".join(
            f"{row['维度名称']}的标准化系数为{row['系数']:.3f}（校正后{significance_text(float(row['BH校正p值']))}）"
            for _, row in significant_scores.iterrows()
        ) + "。"
    add_paragraph(document, score_text)
    add_paragraph(
        document,
        "整体而言，作文分数对五维语言画像的解释力远小于题目组。句法延展维度随分数略有上升，而基础词汇与叙事推进维度略有下降；词汇丰富度并未呈现稳定的线性分数效应。该结果提示，评分可能同时考虑内容切题、组织、准确性等当前宽表尚未覆盖的因素，不能用单一词汇或句法指标替代综合评分。",
    )

    document.add_page_break()
    document.add_heading("7 日本籍与非日本籍的补充分析", level=1)
    add_figure(document, analysis_dir / "figures/10_回归系数森林图.png", 10, "作文分数与日本籍的调整后维度效应")
    japan_regressions = regressions[regressions["模型项"] == "is_japan"].copy()
    add_dataframe_table(
        document,
        japan_regressions,
        columns=["维度名称", "系数", "稳健标准误", "95%CI下限", "95%CI上限", "BH校正p值"],
        widths_dxa=[3360, 1000, 1200, 1200, 1200, 1400],
        font_size=8.3,
    )
    add_table_source(document, "表6  控制篇名代码和作文分数后的日本籍系数；参照组为非日本籍。")
    add_paragraph(
        document,
        "控制题目和分数后，日本籍作文在基础词汇与叙事推进、人称指涉与信息密度维度上更高，在词汇丰富度、句法延展和动作过程维度上更低。该方向与徐勤（2021）关于日本学习者记叙文口语化、句法复杂性及动作描写不足的部分观察存在表面呼应，但两者不能直接等同：本项目的参照组是其他国家学习者而非汉语母语者，而且J2中日本籍样本占118/155，残余题目混杂仍然可能存在。",
    )
    add_callout(
        document,
        "解释限制",
        "本节是控制变量后的探索性补充，不是国籍本质差异结论。国籍样本量不均、教育背景未知、题目分布高度不平衡，均限制了外推。报告不对42个国籍开展多重均值排名。",
        fill="FFF7E8",
    )

    document.add_page_break()
    document.add_heading("8 HSK等级词汇构成", level=1)
    add_figure(document, analysis_dir / "figures/08_HSK等级构成.png", 11, "四组作文的HSK等级词汇构成")
    add_paragraph(
        document,
        "四组作文均以初等HSK词汇为主体，但基础词汇与叙事推进维度显示，等级构成并非独立于写作任务。J2更集中于1至2级词汇和事件推进表达，议论文组则相对增加中等词汇。高等词汇比例在所有组中都较低，因此高等级词的出现次数不适合单独作为写作水平判断；更有解释力的是等级构成、词汇丰富度和篇章功能的联合模式。",
    )

    document.add_page_break()
    document.add_heading("9 与两篇论文的对照", level=1)
    comparison = pd.DataFrame(
        [
            ["词汇丰富度与词汇扩展", "词汇产出丰富度/受限词汇", "高度对应；本项目同时包含MATTR与HSK词种指标"],
            ["基础词汇与叙事推进", "过去叙述、情境表达", "部分对应；本项目更突出基础等级词和趋向表达"],
            ["人称指涉与信息密度", "书面/口语、个人意见", "部分对应；人物题与议论题形成明显两端"],
            ["句法延展与分句复杂度", "语句复杂性", "直接对应；由句长、分句和长句指标构成"],
            ["动作过程与动词链", "动作描写", "直接对应；动词频率和连续动词结构共同加载"],
        ],
        columns=["本项目维度", "论文相近维度", "关系说明"],
    )
    add_dataframe_table(
        document,
        comparison,
        columns=["本项目维度", "论文相近维度", "关系说明"],
        widths_dxa=[2500, 2300, 4560],
        font_size=8.5,
    )
    add_table_source(document, "表7  本项目五维结构与徐勤（2021、2023）多维结果的概念对照。")
    add_paragraph(
        document,
        "三项研究共同出现词汇丰富度、句法复杂度和动作描写维度，说明这些共现模式在汉语二语作文中具有一定重复性。本项目新增的“人称指涉与信息密度”以及“基础词汇与叙事推进”更明显地受到题目构成影响，提示MF/MD维度并非固定能力表，而会随语料体裁、题目、学习者构成和指标体系变化。",
    )
    add_paragraph(
        document,
        "方法上，本报告沿用两篇论文的每千单位标准化、探索性因子分析、Promax旋转、载荷阈值和维度得分思想，同时增加平行分析、变量级MSA、Bootstrap稳定性、稳健标准误和多重检验校正。因而结果在可复现性和不确定性呈现方面更严格，但仍属于探索性而非验证性因子模型。",
    )

    document.add_page_break()
    document.add_heading("10 教学与研究启示", level=1)
    add_paragraph(
        document,
        "第一，作文反馈可以从单项纠错扩展到多维画像。词汇丰富度低并不必然伴随句法延展不足，动作表达密集也不等于篇章逻辑充分。教师可以针对不同维度分别设计词汇替换、人物指涉、复句扩展和动作链描写任务。",
        bold_lead="第一，",
    )
    add_paragraph(
        document,
        "第二，题目效应必须进入评分研究。四个题目组在多维得分上的差异远大于分数的线性效应，若直接合并不同题目建立评分模型，模型可能把题目特征误当成写作能力。后续预测研究应至少纳入题目固定效应，最好在同题目内部交叉验证。",
        bold_lead="第二，",
    )
    add_paragraph(
        document,
        "第三，词汇等级应与功能结合。高等级词比例很低且与题目相关，仅统计高级词数量容易忽略词汇是否服务于叙事推进、人物刻画或论证。HSK等级、词汇丰富度和篇章功能应联合解释。",
        bold_lead="第三，",
    )
    add_paragraph(
        document,
        "第四，定量结果需要回到文本。因子维度给出的是共现方向，不能自动说明具体表达是否准确、自然或切题。匿名例证显示，同一维度的高低端往往对应不同的组织策略，后续应结合原始标注文本增加错误类型和准确性分析。",
        bold_lead="第四，",
    )

    document.add_page_break()
    document.add_heading("11 局限与后续工作", level=1)
    limitations = [
        "没有汉语母语者对照组，不能回答学习者与母语者之间的差异。",
        "题目与体裁部分混杂；J/Y对比只能称为题目与体裁组合差异。",
        "国籍分布与题目高度不均衡，控制变量无法完全替代均衡抽样或匹配设计。",
        "75分样本已排除，各分数段数量不均，分数趋势的连续性受到限制。",
        "PyNLPIR分词、词性和项目词表可能存在自动标注误差，特别是稀有词、专名和多义词。",
        "当前指标主要描述语言形式和资源分布，尚未覆盖切题度、论证质量、错误率和整体连贯性。",
        "因子名称依赖主要载荷的研究者解释，其他语料或指标集可能得到不同维度。",
    ]
    for item in limitations:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.208
        run = paragraph.add_run(item)
        set_run_font(run)
    add_paragraph(
        document,
        "后续最优先的扩展是基于ori_text标注统计错误类型，并构建同题目、同分数、同国籍条件下的匹配比较；其次可在独立样本上进行验证性因子分析，检验五维结构的可迁移性；最后可把维度得分与人工评分维度、错误率和篇章质量共同建模。",
    )

    document.add_page_break()
    document.add_heading("12 结论", level=1)
    add_paragraph(
        document,
        f"基于620篇HSK作文的{column_count}列统计宽表，本报告建立了一个由{selected['feature_count']}项语言指标构成的五维MF/MD结构。该结构在统计诊断与Bootstrap检验中表现稳定，能够将四个题目组的语言差异概括为词汇丰富度、基础词汇与叙事推进、人称指涉与信息密度、句法延展以及动作过程五个相互关联的方向。",
    )
    add_paragraph(
        document,
        "最重要的实证观察是：题目组解释的多维差异明显大于作文分数的线性效应。MF/MD维度适合作为理解作文类型、写作策略和语言资源配置的工具，但不能替代综合评分，也不能在缺少母语者对照和均衡国籍设计时作本质化解释。该分析为后续错误标注、评分建模和跨语料验证提供了可复现基线。",
    )

    document.add_page_break()
    document.add_heading("参考文献", level=1)
    references = [
        "徐勤（2021）．日本人中国語学習者の叙述文における言語的特徴の分析：MF/MD法を使って．《言語文化共同研究プロジェクト》2020，27-41．https://doi.org/10.18910/85007",
        "徐勤（2023）．中国語作文における言語特徴の考察：多次元分析による日本人中国語学習者と中国語母語話者の作文の比較．大阪大学博士学位论文．https://doi.org/10.18910/91825",
        "Biber, D. (1988). Variation Across Speech and Writing. Cambridge University Press.",
    ]
    for item in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(item)
        set_run_font(run, size=10)

    document.add_page_break()
    document.add_heading("附录A 最终因子载荷", level=1)
    loading_columns = [column for column in loadings.columns if column.startswith("D")]
    appendix = loadings[["字段名", "共同度", *loading_columns, "主要维度"]].copy()
    appendix = appendix.sort_values(["主要维度", "共同度"], ascending=[True, False])
    add_dataframe_table(
        document,
        appendix,
        columns=["字段名", "共同度", *loading_columns],
        widths_dxa=[2200, 800, *([1272] * len(loading_columns))],
        font_size=7.2,
    )
    add_table_source(document, "表A1  Promax模式矩阵。报告解释阈值为绝对载荷0.30。")

    document.add_heading("附录B 维度得分敏感性", level=1)
    add_dataframe_table(
        document,
        score_sensitivity,
        columns=["维度名称", "符号求和与回归得分相关"],
        widths_dxa=[6000, 3360],
        font_size=9,
    )
    add_table_source(document, "表B1  论文式符号求和得分与回归型因子得分的Pearson相关。")

    document.add_heading("附录C 可复现性说明", level=1)
    add_paragraph(
        document,
        f"分析固定随机种子为{metadata['seed']}，Horn平行分析使用{metadata['parallel_analysis']['iterations']}次随机矩阵，Bootstrap目标次数为200次。逐字段筛选原因、2至10因子候选诊断、完整载荷、每篇作文维度得分、组间比较和回归结果均保存在《作文多维分析结果.xlsx》中。",
    )
    add_paragraph(
        document,
        "复现顺序为：安装requirements-analysis.txt中的依赖，运行scripts/analyze_composition_mfmd.py生成CSV、JSON与图表，再运行工作簿和报告构建脚本。源宽表、clean_text和论文PDF在全过程中保持只读。",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    build_report(Path(args.analysis_dir).resolve(), Path(args.output).resolve())
    print(f"Report written: {Path(args.output).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
